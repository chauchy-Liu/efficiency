# -*- coding: utf-8 -*-
"""
Created on Wed Oct 19 15:47:43 2022

@author: itadmin
"""
import glob
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH,WD_LINE_SPACING,WD_TAB_ALIGNMENT,WD_UNDERLINE,WD_PARAGRAPH_ALIGNMENT, WD_COLOR_INDEX
from docx.shared import Pt,Cm,Inches, RGBColor
from docx.oxml.ns import qn, nsdecls
import numpy as np
from docx.oxml import OxmlElement, parse_xml
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
import pandas as pd
import math
import os
import io
from db.db import selectFarmInfo, selectAllZuobiaoPicture, selectWindResourceWord, selectAllWindFrequencyPicture, selectAllAirDensityPicture, selectAllTurbulencePicture, selectNavigationBiasDirectionPicture, selectNavigationBiasControlPicture, selectPitchAnglePicture, selectPitchActionPicture, selectTorqueControlPicture, selectWindDirectionPicture, selectPowerCurvePicture, selectCPPicture, selectPitchUnbalancePicture, selectDevicePicture
import algorithms.show_loss_reason_indicator as show_loss_reason_indicator
import algorithms.show_fault_distribute as show_fault_distribute
import algorithms.show_turbine_limit_loss as show_turbine_limit_loss
import algorithms.show_turbine_fault_loss as show_turbine_fault_loss
import algorithms.show_grid_limit_loss as show_grid_limit_loss
import algorithms.show_grid_fault_loss as show_grid_fault_loss
import algorithms.show_technology_loss as show_technology_loss
import algorithms.show_stop_loss as show_stop_loss
import algorithms.show_power_consistence as show_power_consistence
import traceback
from datetime import datetime, timedelta




resourceType_ID = pd.Series([1,2,3,4])
resourceType_name = pd.Series(["一类", "二类", "三类", "四类"])
resourceType = pd.DataFrame({'resourceType_ID':resourceType_ID,'resourceType_name':resourceType_name})

siteType_ID = pd.Series([1,2,3,4])
siteType_name = pd.Series(["高山", "丘陵", "平原", "海上"])
siteType = pd.DataFrame({'siteType_ID':siteType_ID,'siteType_name':siteType_name})

def conver_to_str(data):
    if isinstance(data, float):
        return str(data)
    elif isinstance(data, int):
        return str(data)
    elif isinstance(data, str):
        return data
    
def convert_to_float(data):
    if isinstance(data, float):
        return data
    elif isinstance(data, str):
        return float(data)

def AddFooterNumber(run):
    fldChar1 = OxmlElement('w:fldChar')  # creates a new element
    fldChar1.set(qn('w:fldCharType'), 'begin')  # sets attribute on element
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')  # sets attribute on element
    instrText.text = 'Page'
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'separate')
    t = OxmlElement('w:t')
    t.text = "Seq"
    fldChar2.append(t)
    fldChar4 = OxmlElement('w:fldChar')
    fldChar4.set(qn('w:fldCharType'), 'end')
    r_element = run._r
    r_element.append(fldChar1)
    r_element.append(instrText)
    r_element.append(fldChar2)
    r_element.append(fldChar4)

def InsertPageNumber(Doc):
    footer = Doc.sections[1].footer # 获取第一个节的页脚
    footer.is_linked_to_previous = False  #编号续前一节
    paragraph = footer.paragraphs[0] # 获取页脚的第一个段落
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER#页脚居中对齐
    run_footer=paragraph.add_run() # 添加页脚内容
    AddFooterNumber(run_footer)
    font = run_footer.font
    font.name = 'Times New Roman'#新罗马字体
    font.size = Pt(10)#10号字体
    font.bold = True#加粗


def fill_table(table, x, y, merge_x=-1, merge_y=-1,
               content='', font_name='微软雅黑', font_size=10.5,
               bold=False, italic=False, underline=False, strike=False,
               font_color=None, highlight_color=None,
               paragraph_alignment=None, line_spacing=None,
               width=2.93, height=0.85, vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.CENTER,
               fill_color=None, topbottommargin=None, startendmargin=None,
               picture=None, picture_width=None, picture_height=None):
    """填充表格

    :param table: 要操作的表格
    :param x: 起始行数，从0开始
    :param y: 起始列数，从0开始
    :param merge_x: 要合并的单元格的行数，从0开始
    :param merge_y: 要合并的单元格的列数，从0开始
    :param content: 填充的内容
    :param font_name: 字体名称，如宋体
    :param font_size: 字体大小，单位为 Pt
    :param bold: 字体是否加粗
    :param italic: 字体是否倾斜
    :param underline: 字体是否下划线
    :param strike: 字体是否删除线
    :param font_color: 字体颜色，十六进制字符串，如黑色'000000'
    :param highlight_color: 字体高亮颜色，如黄色取值WD_COLOR_INDEX.YELLOW
    :param paragraph_alignment: 段落水平对齐方式，如居中为WD_PARAGRAPH_ALIGNMENT.CENTER
    :param line_spacing: 段落几倍行距
    :param height: 单元格高度，单位为 Cm
    :param width: 单元格宽度，单位为 Cm
    :param vertical_alignment: 单元格垂直对齐方式，默认垂直居中
    :param fill_color: 单元格填充颜色，十六进制字符串，如白色'FFFFFF'
    :param topbottommargin: 单元格上下边距，单位为 Cm
    :param startendmargin: 单元格左右边距，单位为 Cm
    :param picture: 插入的图片路径
    :param picture_width: 单元格宽度，单位为 Cm
    :param picture_height: 单元格高度，单位为 Cm
    """
    if merge_x < 0:
        cell = table.cell(x, y)
    else:
        cell = table.cell(x, y).merge(table.cell(merge_x, merge_y))  # 合并单元格

    run = cell.paragraphs[0].add_run(str(content))  # 填充内容
    run.font.name = font_name  # 字体名称
    run.font.element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    run.font.size = Pt(font_size)  # 字体大小
    run.bold = bold  # 字体加粗
    run.italic = italic  # 字体倾斜
    run.underline = underline  # 字体下划线
    run.font.strike = strike  # 字体删除线

    if font_color:
        run.font.color.rgb = RGBColor.from_string(font_color)  # 字体颜色
    if highlight_color:
        run.font.highlight_color = highlight_color  # 字体高亮颜色

    if paragraph_alignment:
        cell.paragraphs[0].paragraph_format.alignment = paragraph_alignment  # 段落水平对齐方式
    if line_spacing:
        cell.paragraphs[0].paragraph_format.line_spacing = line_spacing  # 段落几倍行距

    cell.width = Cm(width)  # 单元格宽度
    table.rows[x].height = Cm(height)  # 单元格高度，即设置该行的高度
    cell.vertical_alignment = vertical_alignment  # 单元格垂直对齐方式
    if fill_color:  # 单元格填充颜色
        cell._tc.get_or_add_tcPr().append(parse_xml('<w:shd {} w:fill="{}"/>'.format(nsdecls('w'), fill_color)))

    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    if topbottommargin is not None:  # 单元格上下边距
        topbottommargin *= 567
        for m in ['top', 'bottom']:
            node = OxmlElement('w:{}'.format(m))
            node.set(qn('w:w'), str(topbottommargin))
            node.set(qn('w:type'), 'dxa')
            tcMar.append(node)
        tcPr.append(tcMar)
    if startendmargin is not None:  # 单元格左右边距
        startendmargin *= 567  # 1Cm = 567Twips
        for m in ['start', 'end', 'left', 'right']:  # left和right是WPS的左右边距
            node = OxmlElement('w:{}'.format(m))
            node.set(qn('w:w'), str(startendmargin))
            node.set(qn('w:type'), 'dxa')
            tcMar.append(node)
        tcPr.append(tcMar)

    if picture:
        picture_width = Cm(picture_width) if picture_width else None
        picture_height = Cm(picture_height) if picture_height else None
        run.add_picture(picture, picture_width, picture_height)
    return


def write_word(farmInfo, startTime, endTime, execute_time):
    document = Document()
    
    
    windfarm_erji = farmInfo['company']#WindFarm.loc[0,'二级公司']
    windfarm_name = farmInfo['farm_name'] #WindFarm.loc[0,'风电场名']
    windfarm_postion = farmInfo['address'] #WindFarm.loc[0,'地址']
    windfarm_capi = farmInfo['capacity'] #WindFarm.loc[0,'容量']
    windfarm_amount = farmInfo['turbine_num'] #WindFarm.loc[0,'风机台数']
    windfarm_wttpye = farmInfo['turbine_type'] #WindFarm.loc[0,'机型']
    windfarm_date = farmInfo['operate_time'] #WindFarm.loc[0,'并网日期']
    windfarm_siteType = siteType[(siteType['siteType_ID'] == farmInfo['wind_resource'] )]['siteType_name'].values[0] #WindFarm.loc[0,'风资源']
    windfarm_yunyingzhongxin = farmInfo['rccID'] #WindFarm.loc[0,'生产运营中心']
    windfarm_resourceType = resourceType[(resourceType['resourceType_ID'] == farmInfo['wind_resource'] )]['resourceType_name'].values[0] #WindFarm.loc[0,'风资源']
    trubine_type = len(farmInfo['wtid']) #len( np.unique(Turbine_attr['turbineTypeID']))
    typeNameList = list(farmInfo['wtid'].keys()) #np.unique(Turbine_attr['turbineTypeID']).tolist()
    wtidAllList = []
    # for key, value in farmInfo['wtid'].items():
    #     wtidAllList = wtidAllList + value
    
    #纸张设置
    #section1 = document.add_section()
    document.sections[0].page_width = Cm(21.0)
    document.sections[0].page_height = Cm(29.7)
    document.sections[0].left_margin = Cm(2.5)
    document.sections[0].right_margin = Cm(2.5)
    document.sections[0].top_margin = Cm(2.5)
    document.sections[0].bottom_margin = Cm(2.5)
    #页眉
    header = document.sections[0].header
    header.is_linked_to_previous = False
    paragraph_header = header.paragraphs[0]
    paragraph_header_run = paragraph_header.add_run('\t                                                                     国核信息科技有限公司') 
    paragraph_header_run.font.size = Pt(10)
    paragraph_header_run.font.name = '宋体'
    paragraph_header_run._element.rPr.rFonts.set(qn('w:eastAsia'),u'宋体')
    paragraph_header_run.underline = WD_UNDERLINE.SINGLE
    #paragraph_header.text = ' \t \t国家电投集团数字科技有限公司'
    #paragraph_header.style = document.styles['Header']
    
    #########封面
    paragraph = document.add_paragraph()
    paragraph_run = paragraph.add_run(windfarm_erji + windfarm_name +'能效评估报告')
    paragraph_run.font.size = Pt(24)
    paragraph_run.font.name = '黑体'
    paragraph_run._element.rPr.rFonts.set(qn('w:eastAsia'),u'黑体')
    paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(100)

    paragraph = document.add_paragraph()
    paragraph_run = paragraph.add_run('编制单位：国核信息科技有限公司')
    paragraph_run.font.size = Pt(16)
    paragraph_run.font.name = '仿宋'
    paragraph_run._element.rPr.rFonts.set(qn('w:eastAsia'),u'仿宋')
    paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(300)
    document.add_page_break()

    #####正文
    picture_num = 0
    table_num = 0
    head0 = document.add_paragraph()
    head0.alignment = WD_ALIGN_PARAGRAPH.CENTER
    head0_run = head0.add_run(str(windfarm_erji + windfarm_name +'能效评估报告')) 
    head0_run.font.size = Pt(20)
    head0_run.font.name = '黑体'
    head0_run._element.rPr.rFonts.set(qn('w:eastAsia'),u'黑体')
    head0.paragraph_format.space_after = Pt(10)

    paragraph_text = str(windfarm_erji + windfarm_name+'位于'+str(windfarm_postion)+'，装机容量'+str(windfarm_capi)+'MW，共计'+str(int(float(windfarm_amount)))+'台机组，机组型号：'+windfarm_wttpye+'，并网时间'+str(windfarm_date)+'。'+
                         '风场接入'+str(windfarm_yunyingzhongxin)+'，风场地形为'+str(windfarm_siteType)+'，属'+str(windfarm_resourceType)+'风资源类型'+'。'+
                         '本次发电场能效评估时间段位为'+str(startTime)+'至'+str(endTime)+'，数据源来自集团公司产业数据中台，评估维度包括：风资源分析、损失电量分析、机组控制性能分析等。')
    paragraph = document.add_paragraph()
    paragraph_run = paragraph.add_run(paragraph_text)
    paragraph_run.font.size = Pt(12)
    paragraph_run.font.name = '仿宋'
    paragraph_run._element.rPr.rFonts.set(qn('w:eastAsia'),u'仿宋')
    paragraph.paragraph_format.first_line_indent = paragraph_run.font.size * 2
    paragraph.paragraph_format.line_spacing = Pt(25)
    paragraph.paragraph_format.space_after = Pt(0)

    #小节1
    paragraph0 = document.add_paragraph(style='ListNumber')
    paragraph0_run = paragraph0.add_run('发电情况统计')
    paragraph0_run.font.size = Pt(14)
    paragraph0_run.font.name = '黑体'
    paragraph0_run._element.rPr.rFonts.set(qn('w:eastAsia'),u'黑体')
    paragraph0.paragraph_format.space_before = Pt(12)
    paragraph0.paragraph_format.space_after = Pt(6)
    
    #document.save(path_farm + '/' + windfarm_name + '能效评估报告.docx')
    
    
    #文字
    #trubine_type = 2
    figure_num = str()
    for i in range(trubine_type):
        if i == 0:
            figure_num = figure_num + str('图'+str(int(i+1)))
        else:
            figure_num = figure_num + str('、图'+str(int(i+1)))
            
    paragraph0_1_text = str('风场详细指标数据见下表。各机组发电量及风速统计如'+figure_num+'所示。')   #str('风场检测周期内的异常情况总结如下表，表中的异常情况会对机组的安全性、关键部件寿命及发电性能造成影响。')
    paragraph0_1 = document.add_paragraph()
    paragraph0_1_run = paragraph0_1.add_run(paragraph0_1_text)
    paragraph0_1_run.font.size = Pt(12)
    paragraph0_1_run.font.name = '仿宋'
    paragraph0_1_run._element.rPr.rFonts.set(qn('w:eastAsia'),u'仿宋')
    paragraph0_1.paragraph_format.first_line_indent = paragraph_run.font.size * 2
    paragraph0_1.paragraph_format.line_spacing = Pt(25)
    paragraph0_1.paragraph_format.space_after = Pt(10)
    
    #表例
    table_num = table_num + 1
    paragraph0_2_text = str('表'+str(table_num)+ '风场发电情况统计表')
    paragraph0_2 = document.add_paragraph()
    paragraph0_2_run = paragraph0_2.add_run(paragraph0_2_text)
    paragraph0_2_run.font.size = Pt(10)
    paragraph0_2_run.font.name = '黑体'
    paragraph0_2_run._element.rPr.rFonts.set(qn('w:eastAsia'),u'黑体')
    #paragraph0_2.paragraph_format.first_line_indent = paragraph_run.font.size * 2
    paragraph0_2.paragraph_format.space_after = Pt(0)
    paragraph0_2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    
    
    #表1
    device_num = trubine_type + 1
    column_width = [2.0,1.36,1.45,1.9,1.42,1.50,1.50,1.50,1.65,1.65,1.3,1.3]
    tablez0_2 = document.add_table(device_num,12,style='Table Grid')
    tablez0_2.alignment = WD_TAB_ALIGNMENT.CENTER
    enyPictureList = []
    metric_tongjiList = []
    for i in range(device_num):
        for j in range(12):
            if i == 0:
                if j == 0:
                    fill_table(tablez0_2, x=i, y=j, content='机型', font_name='黑体', font_size=10, bold=False, paragraph_alignment=WD_PARAGRAPH_ALIGNMENT.CENTER, startendmargin=0, width=column_width[j],height=1.5) 
                elif j == 1:
                    fill_table(tablez0_2, x=i, y=j, content='台数', font_name='黑体', font_size=10, bold=False, paragraph_alignment=WD_PARAGRAPH_ALIGNMENT.CENTER, startendmargin=0, width=column_width[j],height=1.5) 
                elif j == 2:
                    fill_table(tablez0_2, x=i, y=j, content='平均风速(m/s)', font_name='黑体', font_size=10, bold=False, paragraph_alignment=WD_PARAGRAPH_ALIGNMENT.CENTER, startendmargin=0, width=column_width[j],height=1.5) 
                elif j == 3:
                    fill_table(tablez0_2, x=i, y=j, content='实发电量(万kWh)', font_name='黑体', font_size=10, bold=False, paragraph_alignment=WD_PARAGRAPH_ALIGNMENT.CENTER, startendmargin=0, width=column_width[j]) 
                elif j == 4:
                    fill_table(tablez0_2, x=i, y=j, content='等效利用小时数(h)', font_name='黑体', font_size=10, bold=False, paragraph_alignment=WD_PARAGRAPH_ALIGNMENT.CENTER, startendmargin=0, width=column_width[j],height=1.5) 
                elif j == 5:
                    fill_table(tablez0_2, x=i, y=j, content='风能可利用率', font_name='黑体', font_size=10, bold=False, paragraph_alignment=WD_PARAGRAPH_ALIGNMENT.CENTER, startendmargin=0, width=column_width[j],height=1.5) 
                elif j == 6:
                    fill_table(tablez0_2, x=i, y=j, content='时间可利用率', font_name='黑体', font_size=10, bold=False, paragraph_alignment=WD_PARAGRAPH_ALIGNMENT.CENTER, startendmargin=0, width=column_width[j],height=1.5)
                elif j == 7:
                    fill_table(tablez0_2, x=i, y=j, content='MTBT(h)', font_name='黑体', font_size=10, bold=False, paragraph_alignment=WD_PARAGRAPH_ALIGNMENT.CENTER, startendmargin=0, width=column_width[j],height=1.5)
                elif j == 8:
                    fill_table(tablez0_2, x=i, y=j, content='损失电量(万kWh)', font_name='黑体', font_size=10, bold=False, paragraph_alignment=WD_PARAGRAPH_ALIGNMENT.CENTER, startendmargin=0, width=column_width[j],height=1.5)
                elif j == 9:
                    fill_table(tablez0_2, x=i, y=j, content='限电率', font_name='黑体', font_size=10, bold=False, paragraph_alignment=WD_PARAGRAPH_ALIGNMENT.CENTER, startendmargin=0, width=column_width[j],height=1.5)
                elif j == 10:
                    fill_table(tablez0_2, x=i, y=j, content='故障恢复时间(h)', font_name='黑体', font_size=10, bold=False, paragraph_alignment=WD_PARAGRAPH_ALIGNMENT.CENTER, startendmargin=0, width=column_width[j],height=1.5)
                elif j == 11:
                    fill_table(tablez0_2, x=i, y=j, content='无故障时间(天)', font_name='黑体', font_size=10, bold=False, paragraph_alignment=WD_PARAGRAPH_ALIGNMENT.CENTER, startendmargin=0, width=column_width[j],height=1.5)
            else:
                # path = str(str(farmInfo['path_farm'])+'/'+str(np.unique(Turbine_attr['turbineTypeID'])[i-1]))
                # metric_tongji = pd.read_csv(str(path+'/tongji.csv'),header=[0],index_col=[0])
                try:
                    if j==0:
                        metric_tongji = show_loss_reason_indicator.analyse(farmInfo['farm_name'], [str(typeNameList[i-1])], startTime, endTime)
                        metric_tongjiList.append(metric_tongji)
                        loss_power = eval(metric_tongji['reason']['turbineFaultLoss']) + eval(metric_tongji['reason']['limGridLoss']) + eval(metric_tongji['reason']['gridFaultLoss']) +eval(metric_tongji['reason']['limTurbineLoss']) + eval(metric_tongji['reason']['technologyLoss']) + eval(metric_tongji['reason']['stopLoss'])
                        enyPictureList.append(metric_tongji['figure'])
                    if j == 0:
                        fill_table(tablez0_2, x=i, y=j, content=str(str(typeNameList[i-1])), font_name='Times New Roman', font_size=10, bold=False, paragraph_alignment=WD_PARAGRAPH_ALIGNMENT.CENTER, startendmargin=0, width=column_width[j]) 
                    elif j == 1:
                        fill_table(tablez0_2, x=i, y=j, content=int(len(farmInfo['wtid'][typeNameList[i-1]])), font_name='Times New Roman', font_size=10, bold=False, paragraph_alignment=WD_PARAGRAPH_ALIGNMENT.CENTER, startendmargin=0, width=column_width[j]) 
                    elif j == 2:
                        fill_table(tablez0_2, x=i, y=j, content='{:.2f}'.format(eval(metric_tongji['indicator']['meanWindSpeed'])), font_name='Times New Roman', font_size=10, bold=False, paragraph_alignment=WD_PARAGRAPH_ALIGNMENT.CENTER, startendmargin=0, width=column_width[j]) 
                    elif j == 3:
                        fill_table(tablez0_2, x=i, y=j, content='{:.1f}'.format(eval(metric_tongji['indicator']['actualPower'])), font_name='Times New Roman', font_size=10, bold=False, paragraph_alignment=WD_PARAGRAPH_ALIGNMENT.CENTER, startendmargin=0, width=column_width[j]) 
                    elif j == 4:
                        fill_table(tablez0_2, x=i, y=j, content='{:.1f}'.format(eval(metric_tongji['indicator']['validHour'])), font_name='Times New Roman', font_size=10, bold=False, paragraph_alignment=WD_PARAGRAPH_ALIGNMENT.CENTER, startendmargin=0, width=column_width[j]) 
                    elif j == 5:
                        fill_table(tablez0_2, x=i, y=j, content=metric_tongji['indicator']['powerRate'], font_name='Times New Roman', font_size=10, bold=False, paragraph_alignment=WD_PARAGRAPH_ALIGNMENT.CENTER, startendmargin=0, width=column_width[j]) 
                    elif j == 6:
                        fill_table(tablez0_2, x=i, y=j, content=metric_tongji['indicator']['timeAvailableRate'], font_name='Times New Roman', font_size=10, bold=False, paragraph_alignment=WD_PARAGRAPH_ALIGNMENT.CENTER, startendmargin=0, width=column_width[j])
                    elif j == 7:
                        fill_table(tablez0_2, x=i, y=j, content='{:.1f}'.format(eval(metric_tongji['indicator']['MTBF'])), font_name='Times New Roman', font_size=10, bold=False, paragraph_alignment=WD_PARAGRAPH_ALIGNMENT.CENTER, startendmargin=0, width=column_width[j])
                    elif j == 8:
                        fill_table(tablez0_2, x=i, y=j, content='{:.1f}'.format(loss_power), font_name='Times New Roman', font_size=10, bold=False, paragraph_alignment=WD_PARAGRAPH_ALIGNMENT.CENTER, startendmargin=0, width=column_width[j])
                    elif j == 9:
                        fill_table(tablez0_2, x=i, y=j, content=metric_tongji['indicator']['limitPowerRate'], font_name='Times New Roman', font_size=10, bold=False, paragraph_alignment=WD_PARAGRAPH_ALIGNMENT.CENTER, startendmargin=0, width=column_width[j])
                    elif j == 10:
                        fill_table(tablez0_2, x=i, y=j, content=eval(metric_tongji['indicator']['faultStoreTime']), font_name='Times New Roman', font_size=10, bold=False, paragraph_alignment=WD_PARAGRAPH_ALIGNMENT.CENTER, startendmargin=0, width=column_width[j])
                    elif j == 11:
                        fill_table(tablez0_2, x=i, y=j, content=eval(metric_tongji['indicator']['noFaultTime']), font_name='Times New Roman', font_size=10, bold=False, paragraph_alignment=WD_PARAGRAPH_ALIGNMENT.CENTER, startendmargin=0, width=column_width[j])
                except Exception as e:

                    errorInfomation = traceback.format_exc()
                    # print("数据库中表eny_wspd_all在时段"+str(startTime)+"到"+str(endTime)+"没有机型"+str(typeNameList[i-1]))
                    print("############################数据库中表eny_wspd_all在时段"+str(startTime)+"到"+str(endTime)+"没有机型"+str(typeNameList[i-1])+'的数据导致报错#################################')
                    print(f'{errorInfomation}')
                    # logger.info(f'\033[31m{errorInfomation}\033[0m')
                    # logger.info(f'\033[33m指标报错：{e}\033[0m')

    #图1
    for i in range(len(metric_tongjiList)): #trubine_type
        # figure_path = str(str(path_farm)+'/'+str(np.unique(Turbine_attr['turbineTypeID'])[i])+'/'+'eny.png')
        wtidAllList += farmInfo['wtid'][typeNameList[i]]
        if i < len(enyPictureList):
            paragraph_picture0_3 = document.add_paragraph()
            paragraph_picture0_3.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph_picture0_3_run = paragraph_picture0_3.add_run()
            paragraph_picture0_3_run.add_picture(enyPictureList[i],width=Cm(14)) #str(figure_path)

            if i == 0:
                paragraph_picture0_3.paragraph_format.space_before = Pt(15)
            paragraph_picture0_3.paragraph_format.space_after = Pt(0)
            #图示
            picture_num += 1
            paragraph0_3_text = str('图'+str(picture_num)+' '+str(typeNameList[i])+'机组实发电量、风速统计图')
            paragraph0_3 = document.add_paragraph()
            paragraph0_3_run = paragraph0_3.add_run(paragraph0_3_text)
            paragraph0_3_run.font.size = Pt(10)
            paragraph0_3_run.font.name = '黑体'
            paragraph0_3_run._element.rPr.rFonts.set(qn('w:eastAsia'),u'黑体')
            #paragraph0_3.paragraph_format.first_line_indent = paragraph_run.font.size * 2
            paragraph0_3.paragraph_format.space_after = Pt(10)
            paragraph0_3.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #文字
    paragraph0_4_text = '如图'+str(picture_num+1)+'为各机组的坐标图，圆点大小表征发电量，颜色表征海拔高度，机组发电量或风速与海拔及机组位置相关性明显or不明显！！！！！'
    paragraph0_4 = document.add_paragraph()
    paragraph0_4_run = paragraph0_4.add_run(paragraph0_4_text)
    paragraph0_4_run.font.size = Pt(12)
    paragraph0_4_run.font.name = '仿宋'
    paragraph0_4_run._element.rPr.rFonts.set(qn('w:eastAsia'),u'仿宋')
    paragraph0_4.paragraph_format.first_line_indent = paragraph_run.font.size * 2
    paragraph0_4.paragraph_format.line_spacing = Pt(25)
    paragraph0_4.paragraph_format.space_after = Pt(10)

    #图2
    figure_path = selectAllZuobiaoPicture(farmInfo['farm_name'], startTime, endTime)#str(str(path_farm)+'/'+'zuobiao_all.png')
    paragraph_picture0_5 = document.add_paragraph()
    paragraph_picture0_5.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph_picture0_5_run = paragraph_picture0_5.add_run()
    paragraph_picture0_5_run.add_picture(figure_path,width=Cm(14))
    paragraph_picture0_5.paragraph_format.space_after = Pt(0)
    
    #图示
    picture_num += 1
    paragraph0_5_text = str('图'+str(picture_num)+' '+'机位分布图')
    paragraph0_5 = document.add_paragraph()
    paragraph0_5_run = paragraph0_5.add_run(paragraph0_5_text)
    paragraph0_5_run.font.size = Pt(10)
    paragraph0_5_run.font.name = '黑体'
    paragraph0_5_run._element.rPr.rFonts.set(qn('w:eastAsia'),u'黑体')
    #paragraph0_5.paragraph_format.first_line_indent = paragraph_run.font.size * 2
    paragraph0_5.paragraph_format.space_after = Pt(10)
    paragraph0_5.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
##########################################################################################

    #小节2
    paragraph1 = document.add_paragraph(style='ListNumber')
    paragraph1_run = paragraph1.add_run('风资源情况分析')
    paragraph1_run.font.size = Pt(14)
    paragraph1_run.font.name = '黑体'
    paragraph1_run._element.rPr.rFonts.set(qn('w:eastAsia'),u'黑体')
    paragraph1.paragraph_format.space_before = Pt(12)
    paragraph1.paragraph_format.space_after = Pt(6)

    #document.save(path_farm + '/' + windfarm_name + '能效评估报告.docx')
    
    #文字
    picture_num += 1
    wind_freq = selectWindResourceWord(farmInfo['farm_name'], startTime, endTime)
    # 将freq列转换为数值类型
    wind_freq['freq'] = pd.to_numeric(wind_freq['freq'], errors='coerce')
    top6_freq = wind_freq.nlargest(6, 'freq', keep='all')
    min_speed = top6_freq['windbin'].min()
    max_speed = top6_freq['windbin'].max()
    paragraph1_1_text = '风电场评估周期内平均风速'+'{:.2f}'.format(wind_freq.iloc[0]['wind_mean'])+'m/s，最大风速'+'{:.2f}'.format(wind_freq.iloc[0]['wind_max'])+'m/s，风频主要集中在'+'{:.2f}'.format(min_speed)+'～'+'{:.2f}'.format(max_speed)+'m/s风速区间，图'+str(picture_num)+'为风频分布图。' 
    paragraph1_1 = document.add_paragraph()
    paragraph1_1_run = paragraph1_1.add_run(paragraph1_1_text)
    paragraph1_1_run.font.size = Pt(12)
    paragraph1_1_run.font.name = '仿宋'
    paragraph1_1_run._element.rPr.rFonts.set(qn('w:eastAsia'),u'仿宋')
    paragraph1_1.paragraph_format.first_line_indent = paragraph_run.font.size * 2
    paragraph1_1.paragraph_format.line_spacing = Pt(25)
    paragraph1_1.paragraph_format.space_after = Pt(12)

    #图3
    figure_path = selectAllWindFrequencyPicture(farmInfo['farm_name'], startTime, endTime)#str(str(path_farm)+'/'+'windfreq.png')
    paragraph_picture1_2 = document.add_paragraph()
    paragraph_picture1_2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph_picture1_2_run = paragraph_picture1_2.add_run()
    paragraph_picture1_2_run.add_picture(figure_path,width=Cm(14))
    paragraph_picture1_2.paragraph_format.space_after = Pt(0)
    #图示
    paragraph1_2_text = str('图'+str(picture_num)+' '+'风频分布图')
    paragraph1_2 = document.add_paragraph()
    paragraph1_2_run = paragraph1_2.add_run(paragraph1_2_text)
    paragraph1_2_run.font.size = Pt(10)
    paragraph1_2_run.font.name = '黑体'
    paragraph1_2_run._element.rPr.rFonts.set(qn('w:eastAsia'),u'黑体')
    #paragraph1_2.paragraph_format.first_line_indent = paragraph_run.font.size * 2
    paragraph1_2.paragraph_format.space_after = Pt(10)
    paragraph1_2.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #文字
    picture_num += 1
    mean_rho = wind_freq.iloc[0]['mean_rho'] #month_data['rho'].mean()
    max_speed_month = wind_freq.iloc[0]['max_speed_month'] #month_data['wspd'].idxmax()
    paragraph1_3_text = '图'+str(picture_num)+'为统计时间段内风场风速、空气密度变化图（红色代表空气密度、蓝色代表风速），平均空气密度'+'{:.2f}'.format(mean_rho)+'kg/m'+chr(179)+'，'+str(max_speed_month)+'份风速相对较大。'
    paragraph1_3 = document.add_paragraph()
    paragraph1_3_run = paragraph1_3.add_run(paragraph1_3_text)
    paragraph1_3_run.font.size = Pt(12)
    paragraph1_3_run.font.name = '仿宋'
    paragraph1_3_run._element.rPr.rFonts.set(qn('w:eastAsia'),u'仿宋')
    paragraph1_3.paragraph_format.first_line_indent = paragraph_run.font.size * 2
    paragraph1_3.paragraph_format.line_spacing = Pt(25)
    paragraph1_3.paragraph_format.space_after = Pt(12)

    #图4
    figure_path = selectAllAirDensityPicture(farmInfo['farm_name'], startTime, endTime)#str(str(path_farm)+'/'+'month.png')
    paragraph_picture1_4 = document.add_paragraph()
    paragraph_picture1_4.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph_picture1_4_run = paragraph_picture1_4.add_run()
    paragraph_picture1_4_run.add_picture(figure_path,width=Cm(14))
    paragraph_picture1_4.paragraph_format.space_after = Pt(0)
    #图示
    
    paragraph1_4_text = str('图'+str(picture_num)+' '+'月度风速、空气密度变化图')
    paragraph1_4 = document.add_paragraph()
    paragraph1_4_run = paragraph1_4.add_run(paragraph1_4_text)
    paragraph1_4_run.font.size = Pt(10)
    paragraph1_4_run.font.name = '黑体'
    paragraph1_4_run._element.rPr.rFonts.set(qn('w:eastAsia'),u'黑体')
    paragraph1_4.paragraph_format.space_after = Pt(6)
    paragraph1_4.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #文字
    paragraph1_5_text = '因为没有风向及机舱真北标定的相关数据，因此假定各机组偏航0度位置为正北方向，各机组风向玫瑰图如下图所示，大部分机组的主风能方向相对集中。'
    paragraph1_5 = document.add_paragraph()
    paragraph1_5_run = paragraph1_5.add_run(paragraph1_5_text)
    paragraph1_5_run.font.size = Pt(12)
    paragraph1_5_run.font.name = '仿宋'
    paragraph1_5_run._element.rPr.rFonts.set(qn('w:eastAsia'),u'仿宋')
    paragraph1_5.paragraph_format.first_line_indent = paragraph_run.font.size * 2
    paragraph1_5.paragraph_format.line_spacing = Pt(25)
    paragraph1_5.paragraph_format.space_after = Pt(6)

    #图5    
    row_num = math.ceil(len(wtidAllList)/2)
    paragraph_picture1_5 = document.add_paragraph()
    paragraph_picture1_5.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # paragraph_picture1_5_run = paragraph_picture1_5.add_run()
    # paragraph_picture1_5_run.add_picture(str(figure_rose),width=Cm(14))
    paragraph_picture1_5.paragraph_format.space_after = Pt(0)

    tablez1_5 = document.add_table(row_num,2)
    #tablez.autofit = True
    tablez1_5.alignment = WD_TAB_ALIGNMENT.CENTER
    trubine_type_i = 0
    num = 0
    # path = str(str(path_farm)+'/'+str(np.unique(Turbine_attr['turbineTypeID'])[trubine_type_i]))
    wtids = farmInfo['wtid'][typeNameList[trubine_type_i]]#Turbine_attr[Turbine_attr['turbineTypeID']==np.unique(Turbine_attr['turbineTypeID'])[trubine_type_i]]['name']
    turbine_num = len(wtids)
    for i in range(row_num):
        for j in range(2):
            if ((i*2+j+1) > turbine_num):
                if (i*2+j+1)>len(wtidAllList) or trubine_type_i >= len(typeNameList)-1:
                    break
                trubine_type_i = trubine_type_i + 1
                num = 0
                # path = str(str(path_farm)+'/'+str(np.unique(Turbine_attr['turbineTypeID'])[trubine_type_i]))
                wtids = farmInfo['wtid'][typeNameList[trubine_type_i]] #Turbine_attr[Turbine_attr['turbineTypeID']==np.unique(Turbine_attr['turbineTypeID'])[trubine_type_i]]['name']
                turbine_num = turbine_num + len(wtids)
            figure_path = selectWindDirectionPicture(farmInfo['farm_name'], typeNameList[trubine_type_i], wtids[num], startTime, endTime)#str(path+'/'+str(wtids.iloc[num])+'风向玫瑰图.png')
            if figure_path != None: #os.path.exists(figure_path)==True:
                fill_table(tablez1_5, x=i, y=j, width=8.0, height=7.0,startendmargin=0, picture=figure_path, picture_width=8, picture_height=7.12,paragraph_alignment=WD_PARAGRAPH_ALIGNMENT.CENTER)
                num = num + 1
            
    #表放的位置
    paragraph_picture1_5._p.addnext(tablez1_5._tbl)

    #图示
    picture_num += 1
    paragraph1_5_text = str('图'+str(picture_num)+' '+'各机组风向玫瑰图')
    paragraph1_5 = document.add_paragraph()
    paragraph1_5_run = paragraph1_5.add_run(paragraph1_5_text)
    paragraph1_5_run.font.size = Pt(10)
    paragraph1_5_run.font.name = '黑体'
    paragraph1_5_run._element.rPr.rFonts.set(qn('w:eastAsia'),u'黑体')
    paragraph1_5.paragraph_format.space_after = Pt(6)
    paragraph1_5.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    #文字
    if wind_freq.iloc[0]['turbulence_flag15'] > 0:#(np.nanmax(wind_ti_alltype['windbin']) >= 15.0)&(15.0 in wind_ti_alltype['windbin'].values):
        ti_windbin15 =  wind_freq.iloc[0]['turbulence'] #round(wind_ti_alltype[wind_ti_alltype['windbin']==15.0]['ti'].values[0],2)
        if ti_windbin15 > 0.12:
            str_temp = '，湍流强度较大。'
        else:
            str_temp = '，湍流强度较小。'
        paragraph1_6_text = ('湍流强度是描述风速随时间变化程度的参数，根据机组风速仪测风数据计算，全场机组在轮毂高度处各风速段的平均湍流强度变化曲线见下图。15m/s风速段湍流强度为'+
                             str(ti_windbin15)+str_temp)
        paragraph1_6 = document.add_paragraph()
        paragraph1_6_run = paragraph1_6.add_run(paragraph1_6_text)
        paragraph1_6_run.font.size = Pt(12)
        paragraph1_6_run.font.name = '仿宋'
        paragraph1_6_run._element.rPr.rFonts.set(qn('w:eastAsia'),u'仿宋')
        paragraph1_6.paragraph_format.first_line_indent = paragraph_run.font.size * 2
        paragraph1_6.paragraph_format.line_spacing = Pt(25)
        paragraph1_6.paragraph_format.space_after = Pt(6)
    else:
        if wind_freq.iloc[0]['turbulence']: #np.nanmean(wind_ti_alltype['ti']) > 0.12:
            str_temp = '湍流强度较大。'
        else:
            str_temp = '湍流强度较小。'
        paragraph1_6_text = ('湍流强度是描述风速随时间变化程度的参数，根据机组风速仪测风数据计算，全场机组在轮毂高度处各风速段的平均湍流强度变化曲线见下图，风场'+str_temp)
        paragraph1_6 = document.add_paragraph()
        paragraph1_6_run = paragraph1_6.add_run(paragraph1_6_text)
        paragraph1_6_run.font.size = Pt(12)
        paragraph1_6_run.font.name = '仿宋'
        paragraph1_6_run._element.rPr.rFonts.set(qn('w:eastAsia'),u'仿宋')
        paragraph1_6.paragraph_format.first_line_indent = paragraph_run.font.size * 2
        paragraph1_6.paragraph_format.line_spacing = Pt(25)
        paragraph1_6.paragraph_format.space_after = Pt(6)
        
    
    #图4
    picture_num += 1
    figure_path = selectAllTurbulencePicture(farmInfo['farm_name'], startTime, endTime) #str(str(path_farm)+'/'+'湍流曲线.png')
    paragraph_picture1_4 = document.add_paragraph()
    paragraph_picture1_4.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph_picture1_4_run = paragraph_picture1_4.add_run()
    paragraph_picture1_4_run.add_picture(figure_path,width=Cm(14))
    paragraph_picture1_4.paragraph_format.space_after = Pt(0)
    #图示
    
    paragraph1_4_text = str('图'+str(picture_num)+' '+'各风速段湍流强度变化曲线图')
    paragraph1_4 = document.add_paragraph()
    paragraph1_4_run = paragraph1_4.add_run(paragraph1_4_text)
    paragraph1_4_run.font.size = Pt(10)
    paragraph1_4_run.font.name = '黑体'
    paragraph1_4_run._element.rPr.rFonts.set(qn('w:eastAsia'),u'黑体')
    paragraph1_4.paragraph_format.space_after = Pt(6)
    paragraph1_4.alignment = WD_ALIGN_PARAGRAPH.CENTER

    
##########################################################################################

    #小节3
    paragraph2 = document.add_paragraph(style='ListNumber')
    paragraph2_run = paragraph2.add_run('损失电量分析')
    paragraph2_run.font.size = Pt(14)
    paragraph2_run.font.name = '黑体'
    paragraph2_run._element.rPr.rFonts.set(qn('w:eastAsia'),u'黑体')
    paragraph2.paragraph_format.space_before = Pt(12)
    paragraph2.paragraph_format.space_after = Pt(6)
    
    
    #文字
    for i in range(len(metric_tongjiList)): #trubine_type
        # path = str(str(path_farm)+'/'+str(np.unique(Turbine_attr['turbineTypeID'])[i]))
        metric_tongji = metric_tongjiList[i]#pd.read_csv(str(path+'/tongji.csv'),header=[0],index_col=[0])
        loss_power = eval(metric_tongji['reason']['turbineFaultLoss']) + eval(metric_tongji['reason']['limGridLoss']) + eval(metric_tongji['reason']['gridFaultLoss']) +eval(metric_tongji['reason']['limTurbineLoss']) + eval(metric_tongji['reason']['technologyLoss']) + eval(metric_tongji['reason']['stopLoss'])
        paragraph2_1_text = ('分析时段内，'+windfarm_name+str(typeNameList[i])+'机型各类损失电量共计'+'{:.1f}'.format(loss_power)+'kWh，占应发电量'+'{:.2f}'.format(100-float(eval(metric_tongji['indicator']['powerRate'].strip('%'))))+'%。'+
                             '其中风机机组故障损失电量'+'{:.1f}'.format(eval(metric_tongji['reason']['turbineFaultLoss'])*10000)+'kWh，'+
                             '电网限电损失电量'+'{:.1f}'.format(eval(metric_tongji['reason']['limGridLoss'])*10000)+'kWh，'+
                             '电网故障损失电量'+'{:.1f}'.format(eval(metric_tongji['reason']['gridFaultLoss'])*10000)+'kWh，'+
                             '机组自限电损失电量'+'{:.1f}'.format(eval(metric_tongji['reason']['limTurbineLoss'])*10000)+'kWh，'+
                             '机组技术待机损失电量'+'{:.1f}'.format(eval(metric_tongji['reason']['technologyLoss'])*10000)+'kWh，'+
                             '计划停机损失电量'+'{:.1f}'.format(eval(metric_tongji['reason']['stopLoss'])*10000)+'kWh。')

        paragraph2_1 = document.add_paragraph()
        paragraph2_1_run = paragraph2_1.add_run(paragraph2_1_text)
        paragraph2_1_run.font.size = Pt(12)
        paragraph2_1_run.font.name = '仿宋'
        paragraph2_1_run._element.rPr.rFonts.set(qn('w:eastAsia'),u'仿宋')
        paragraph2_1.paragraph_format.first_line_indent = paragraph_run.font.size * 2
        paragraph2_1.paragraph_format.line_spacing = Pt(25)
        paragraph2_1.paragraph_format.space_after = Pt(6)
    
    paragraph2_2_text = '3.1. 机组故障损失 '
    paragraph2_2 = document.add_paragraph()
    paragraph2_2_run = paragraph2_2.add_run(paragraph2_2_text)
    paragraph2_2_run.font.size = Pt(12)
    paragraph2_2_run.font.name = '黑体'
    paragraph2_2_run._element.rPr.rFonts.set(qn('w:eastAsia'),u'黑体')
    # paragraph2_2.paragraph_format.first_line_indent = paragraph_run.font.size * 2
    paragraph2_2.paragraph_format.space_before = Pt(12) 
    paragraph2_2.paragraph_format.space_after = Pt(0) 
    
    for typei in range(len(metric_tongjiList)):  #trubine_type
        # path = str(str(path_farm)+'/'+str(np.unique(Turbine_attr['turbineTypeID'])[typei]))
        metric_tongji = metric_tongjiList[typei]#pd.read_csv(str(path+'/tongji.csv'),header=[0],index_col=[0])
        
        # path = str(str(path_farm)+'/'+str(np.unique(Turbine_attr['turbineTypeID'])[typei]))
        fault_loss_all = show_turbine_fault_loss.analyse(farmInfo['farm_name'], [typeNameList[typei]], startTime, endTime)#['table']#pd.read_csv(str(path+'/fault_loss_all.csv'),header=[0],index_col=[0])
        if len(fault_loss_all) > 0:
            fault_loss_all = fault_loss_all['table']
        
        if (eval(metric_tongji['reason']['turbineFaultLoss'])*10000 > 200 and len(fault_loss_all)>0): #&(len(fault_loss_all)>0
            paragraph2_1_text = ('分析时段内'+str(typeNameList[typei])+'机型各机组故障损失电量共'+'{:.1f}'.format(eval(metric_tongji['reason']['turbineFaultLoss'])*10000)+'kWh，折合等效小时'+
                                '{:.1f}'.format(eval(metric_tongji['reason']['turbineFaultLoss'])*10000/windfarm_capi/1000)+'h，故障频次及详细故障损失统计见下图表，其中***故障发生频次较高。')   
            paragraph2_1 = document.add_paragraph()
            paragraph2_1_run = paragraph2_1.add_run(paragraph2_1_text)
            paragraph2_1_run.font.size = Pt(12)
            paragraph2_1_run.font.name = '仿宋'
            paragraph2_1_run._element.rPr.rFonts.set(qn('w:eastAsia'),u'仿宋')
            paragraph2_1.paragraph_format.first_line_indent = paragraph_run.font.size * 2
            paragraph2_1.paragraph_format.line_spacing = Pt(25)
            paragraph2_1.paragraph_format.space_after = Pt(6)
            paragraph2_1.paragraph_format.space_before = Pt(6)
            #图补
            _, figure_path = show_fault_distribute.analyse(farmInfo['farm_name'], [typeNameList[typei]], startTime, endTime)#str(str(path_farm)+'/'+str(np.unique(Turbine_attr['turbineTypeID'])[typei])+'/'+'fault.png')
            paragraph_picture2_4 = document.add_paragraph()
            paragraph_picture2_4.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph_picture2_4_run = paragraph_picture2_4.add_run()
            paragraph_picture2_4_run.add_picture(figure_path,width=Cm(14))
            paragraph_picture2_4.paragraph_format.space_after = Pt(0)
            #图示
            picture_num += 1
            paragraph2_4_text = str('图'+str(picture_num)+' '+str(typeNameList[typei])+'机型故障频次饼状图')
            paragraph2_4 = document.add_paragraph()
            paragraph2_4_run = paragraph2_4.add_run(paragraph2_4_text)
            paragraph2_4_run.font.size = Pt(10)
            paragraph2_4_run.font.name = '黑体'
            paragraph2_4_run._element.rPr.rFonts.set(qn('w:eastAsia'),u'黑体')
            paragraph2_4.paragraph_format.space_after = Pt(16)
            paragraph2_4.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            #表例
            table_num = table_num + 1
            paragraph2_5_text = str('表'+str(table_num)+' '+str(typeNameList[typei])+'机型故障损失统计表')
            paragraph2_5 = document.add_paragraph()
            paragraph2_5_run = paragraph2_5.add_run(paragraph2_5_text)
            paragraph2_5_run.font.size = Pt(10)
            paragraph2_5_run.font.name = '黑体'
            paragraph2_5_run._element.rPr.rFonts.set(qn('w:eastAsia'),u'黑体')
            paragraph2_5.paragraph_format.space_after = Pt(0)
            paragraph2_5.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            #表2
            column_width = [1.25,1.70,1.72,2.35,2.67,2.33,3.99]
            device_num = len(fault_loss_all)+1
            tablez2_5 = document.add_table(device_num,len(column_width),style='Table Grid')
            tablez2_5.alignment = WD_TAB_ALIGNMENT.CENTER
            for i in range(device_num):
                for j in range(len(column_width)):
                    if i == 0:
                        if j == 0:
                            fill_table(tablez2_5, x=i, y=j, content='机位号', font_name='黑体', font_size=10, bold=False, paragraph_alignment=WD_PARAGRAPH_ALIGNMENT.CENTER, startendmargin=0, width=column_width[j]) 
                        elif j == 1:
                            fill_table(tablez2_5, x=i, y=j, content='故障码', font_name='黑体', font_size=10, bold=False, paragraph_alignment=WD_PARAGRAPH_ALIGNMENT.CENTER, startendmargin=0, width=column_width[j]) 
                        elif j == 2:
                            fill_table(tablez2_5, x=i, y=j, content='故障频次', font_name='黑体', font_size=10, bold=False, paragraph_alignment=WD_PARAGRAPH_ALIGNMENT.CENTER, startendmargin=0, width=column_width[j]) 
                        elif j == 3:
                            fill_table(tablez2_5, x=i, y=j, content='故障时长(h)', font_name='黑体', font_size=10, bold=False, paragraph_alignment=WD_PARAGRAPH_ALIGNMENT.CENTER, startendmargin=0, width=column_width[j]) 
                        elif j == 4:
                            fill_table(tablez2_5, x=i, y=j, content='故障损失电量(kWh)', font_name='黑体', font_size=10, bold=False, paragraph_alignment=WD_PARAGRAPH_ALIGNMENT.CENTER, startendmargin=0, width=column_width[j]) 
                        elif j == 5:
                            fill_table(tablez2_5, x=i, y=j, content='故障时平均风速(m/s)', font_name='黑体', font_size=10, bold=False, paragraph_alignment=WD_PARAGRAPH_ALIGNMENT.CENTER, startendmargin=0, width=column_width[j]) 
                        elif j == 6:
                            fill_table(tablez2_5, x=i, y=j, content='故障描述', font_name='黑体', font_size=10, bold=False, paragraph_alignment=WD_PARAGRAPH_ALIGNMENT.CENTER, startendmargin=0, width=column_width[j]) 
                    else:
                        if j == 0:
                            fill_table(tablez2_5, x=i, y=j, content=conver_to_str(fault_loss_all[i-1]['wtid']), font_name='Times New Roman', font_size=10, bold=False, paragraph_alignment=WD_PARAGRAPH_ALIGNMENT.CENTER, startendmargin=0, width=column_width[j]) 
                        elif j == 1:
                            fill_table(tablez2_5, x=i, y=j, content=int(fault_loss_all[i-1]['faultCode']), font_name='Times New Roman', font_size=10, bold=False, paragraph_alignment=WD_PARAGRAPH_ALIGNMENT.CENTER, startendmargin=0, width=column_width[j]) 
                        elif j == 2:
                            fill_table(tablez2_5, x=i, y=j, content=int(fault_loss_all[i-1]['faultCount']), font_name='Times New Roman', font_size=10, bold=False, paragraph_alignment=WD_PARAGRAPH_ALIGNMENT.CENTER, startendmargin=0, width=column_width[j]) 
                        elif j == 3:
                            fill_table(tablez2_5, x=i, y=j, content='{:.2f}'.format(eval(fault_loss_all[i-1]['faultTime'])), font_name='Times New Roman', font_size=10, bold=False, paragraph_alignment=WD_PARAGRAPH_ALIGNMENT.CENTER, startendmargin=0, width=column_width[j]) 
                        elif j == 4:
                            fill_table(tablez2_5, x=i, y=j, content='{:.2f}'.format(eval(fault_loss_all[i-1]['faultLoss'])), font_name='Times New Roman', font_size=10, bold=False, paragraph_alignment=WD_PARAGRAPH_ALIGNMENT.CENTER, startendmargin=0, width=column_width[j]) 
                        elif j == 5:
                            fill_table(tablez2_5, x=i, y=j, content='{:.2f}'.format(eval(fault_loss_all[i-1]['meanWindSpeed'])), font_name='Times New Roman', font_size=10, bold=False, paragraph_alignment=WD_PARAGRAPH_ALIGNMENT.CENTER, startendmargin=0, width=column_width[j]) 
                        elif j == 6:
                            fill_table(tablez2_5, x=i, y=j, content=str(fault_loss_all[i-1]['faultDescribe']), font_name='仿宋', font_size=10, bold=False, paragraph_alignment=WD_PARAGRAPH_ALIGNMENT.CENTER, startendmargin=0, width=column_width[j]) 
        else:
            paragraph2_1_text = '分析时段内'+str(typeNameList[typei])+'机型各机组运行良好，故障极少，损失电量不到200kWh。'      
            paragraph2_1 = document.add_paragraph()
            paragraph2_1_run = paragraph2_1.add_run(paragraph2_1_text)
            paragraph2_1_run.font.size = Pt(12)
            paragraph2_1_run.font.name = '仿宋'
            paragraph2_1_run._element.rPr.rFonts.set(qn('w:eastAsia'),u'仿宋')
            paragraph2_1.paragraph_format.first_line_indent = paragraph_run.font.size * 2
            paragraph2_1.paragraph_format.line_spacing = Pt(25)
            paragraph2_1.paragraph_format.space_after = Pt(6)
            paragraph2_1.paragraph_format.space_before = Pt(6)
    #---------------------------------------------------------------------------------------
    
    paragraph2_2_text = '3.2. 电网限电损失 '
    paragraph2_2 = document.add_paragraph()
    paragraph2_2_run = paragraph2_2.add_run(paragraph2_2_text)
    paragraph2_2_run.font.size = Pt(12)
    paragraph2_2_run.font.name = '黑体'
    paragraph2_2_run._element.rPr.rFonts.set(qn('w:eastAsia'),u'黑体')
    # paragraph2_2.paragraph_format.first_line_indent = paragraph_run.font.size * 2
    paragraph2_2.paragraph_format.space_before = Pt(12)
    paragraph2_2.paragraph_format.space_after = Pt(0) 
    
    for typei in range(len(metric_tongjiList)): # trubine_type
        # path = str(str(path_farm)+'/'+str(np.unique(Turbine_attr['turbineTypeID'])[typei]))
        metric_tongji = metric_tongjiList[i]#pd.read_csv(str(path+'/tongji.csv'),header=[0],index_col=[0])
        
        # path = str(str(path_farm)+'/'+str(np.unique(Turbine_attr['turbineTypeID'])[typei]))
        limgrid_loss_all = show_grid_limit_loss.analyse(farmInfo['farm_name'], [typeNameList[typei]], startTime, endTime)#pd.read_csv(str(path+'/limgrid_loss_all.csv'),header=[0],index_col=[0])
        if len(limgrid_loss_all) > 0:
            limgrid_loss_all = limgrid_loss_all['table']
        
        if (eval(metric_tongji['reason']['limGridLoss'])*10000 > 200) and (len(limgrid_loss_all)>0):
            paragraph2_1_text = ('分析时段内'+str(typeNameList[typei])+'机型各机组电网限电损失电量共'+'{:.1f}'.format(eval(metric_tongji['reason']['limGridLoss'])*10000)+'kWh，折合等效小时'+
                                '{:.1f}'.format(eval(metric_tongji['reason']['limGridLoss'])*10000/eval(windfarm_capi)/1000)+'h，各机组详细统计见下表。')
            paragraph2_1 = document.add_paragraph()
            paragraph2_1_run = paragraph2_1.add_run(paragraph2_1_text)
            paragraph2_1_run.font.size = Pt(12)
            paragraph2_1_run.font.name = '仿宋'
            paragraph2_1_run._element.rPr.rFonts.set(qn('w:eastAsia'),u'仿宋')
            paragraph2_1.paragraph_format.first_line_indent = paragraph_run.font.size * 2
            paragraph2_1.paragraph_format.line_spacing = Pt(25)
            paragraph2_1.paragraph_format.space_after = Pt(6)
            paragraph2_1.paragraph_format.space_before = Pt(6)
            
            #表例
            table_num = table_num + 1
            paragraph2_5_text = str('表'+str(table_num)+' '+str(typeNameList[typei])+'机型电网限电损失统计表')
            paragraph2_5 = document.add_paragraph()
            paragraph2_5_run = paragraph2_5.add_run(paragraph2_5_text)
            paragraph2_5_run.font.size = Pt(10)
            paragraph2_5_run.font.name = '黑体'
            paragraph2_5_run._element.rPr.rFonts.set(qn('w:eastAsia'),u'黑体')
            #paragraph2_5.paragraph_format.first_line_indent = paragraph_run.font.size * 2
            paragraph2_5.paragraph_format.space_after = Pt(0)
            paragraph2_5.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            #表2
            column_width = [1.9,2.7,3.5,2.7]
            device_num = len(limgrid_loss_all)+1
            tablez2_5 = document.add_table(device_num,len(column_width),style='Table Grid')
            tablez2_5.alignment = WD_TAB_ALIGNMENT.CENTER
            for i in range(device_num):
                for j in range(len(column_width)):
                    if i == 0:
                        if j == 0:
                            fill_table(tablez2_5, x=i, y=j, content='机位号', font_name='黑体', font_size=10, bold=False, paragraph_alignment=WD_PARAGRAPH_ALIGNMENT.CENTER, startendmargin=0, width=column_width[j]) 
                        elif j == 1:
                            fill_table(tablez2_5, x=i, y=j, content='限电时长(h)', font_name='黑体', font_size=10, bold=False, paragraph_alignment=WD_PARAGRAPH_ALIGNMENT.CENTER, startendmargin=0, width=column_width[j]) 
                        elif j == 2:
                            fill_table(tablez2_5, x=i, y=j, content='限电损失电量(kWh)', font_name='黑体', font_size=10, bold=False, paragraph_alignment=WD_PARAGRAPH_ALIGNMENT.CENTER, startendmargin=0, width=column_width[j]) 
                        elif j == 3:
                            fill_table(tablez2_5, x=i, y=j, content='限电时平均风速(m/s)', font_name='黑体', font_size=10, bold=False, paragraph_alignment=WD_PARAGRAPH_ALIGNMENT.CENTER, startendmargin=0, width=column_width[j]) 
                    else:
                        if j == 0:
                            fill_table(tablez2_5, x=i, y=j, content=conver_to_str(limgrid_loss_all[i-1]['wtid']), font_name='Times New Roman', font_size=10, bold=False, paragraph_alignment=WD_PARAGRAPH_ALIGNMENT.CENTER, startendmargin=0, width=column_width[j]) 
                        elif j == 1:
                            fill_table(tablez2_5, x=i, y=j, content='{:.2f}'.format(eval(limgrid_loss_all[i-1]['faultTime'])), font_name='Times New Roman', font_size=10, bold=False, paragraph_alignment=WD_PARAGRAPH_ALIGNMENT.CENTER, startendmargin=0, width=column_width[j]) 
                        elif j == 2:
                            fill_table(tablez2_5, x=i, y=j, content='{:.2f}'.format(eval(limgrid_loss_all[i-1]['faultLoss'])), font_name='Times New Roman', font_size=10, bold=False, paragraph_alignment=WD_PARAGRAPH_ALIGNMENT.CENTER, startendmargin=0, width=column_width[j]) 
                        elif j == 3:
                            fill_table(tablez2_5, x=i, y=j, content='{:.2f}'.format(eval(limgrid_loss_all[i-1]['meanWindSpeed'])), font_name='Times New Roman', font_size=10, bold=False, paragraph_alignment=WD_PARAGRAPH_ALIGNMENT.CENTER, startendmargin=0, width=column_width[j]) 
                        
        else:
            paragraph2_1_text = '分析时段内依据产业数据中台数据统计，'+str(typeNameList[typei])+'机型各机组未发生电网限电情况。'      
            paragraph2_1 = document.add_paragraph()
            paragraph2_1_run = paragraph2_1.add_run(paragraph2_1_text)
            paragraph2_1_run.font.size = Pt(12)
            paragraph2_1_run.font.name = '仿宋'
            paragraph2_1_run._element.rPr.rFonts.set(qn('w:eastAsia'),u'仿宋')
            paragraph2_1.paragraph_format.first_line_indent = paragraph_run.font.size * 2
            paragraph2_1.paragraph_format.line_spacing = Pt(25)
            paragraph2_1.paragraph_format.space_after = Pt(6)
            paragraph2_1.paragraph_format.space_before = Pt(6)
    #---------------------------------------------------------------------------------------
    

   #文字
    paragraph2_6_text = '3.3. 机组自限电损失 '
    paragraph2_6 = document.add_paragraph()
    paragraph2_6_run = paragraph2_6.add_run(paragraph2_6_text)
    paragraph2_6_run.font.size = Pt(12)
    paragraph2_6_run.font.name = '黑体'
    paragraph2_6_run._element.rPr.rFonts.set(qn('w:eastAsia'),u'黑体')
    # paragraph6_2.paragraph_format.first_line_indent = paragraph_run.font.size * 2
    paragraph2_6.paragraph_format.space_after = Pt(0) 
    paragraph2_6.paragraph_format.space_before = Pt(12)
    
    for typei in range(len(metric_tongjiList)): # trubine_type
        # path = str(str(path_farm)+'/'+str(np.unique(Turbine_attr['turbineTypeID'])[typei]))
        metric_tongji = metric_tongjiList[typei]#pd.read_csv(str(path+'/tongji.csv'),header=[0],index_col=[0])
        
        # path = str(str(path_farm)+'/'+str(np.unique(Turbine_attr['turbineTypeID'])[typei]))
        limturbine_loss_all = show_turbine_limit_loss.analyse(farmInfo['farm_name'], [typeNameList[typei]], startTime, endTime)#['table']#pd.read_csv(str(path+'/limturbine_loss_all.csv'),header=[0],index_col=[0])
        if len(limturbine_loss_all) > 0:
            limturbine_loss_all = limturbine_loss_all['table']
        
        if (eval(metric_tongji['reason']['limTurbineLoss'])*10000 > 0) and (len(limturbine_loss_all)>0):
            paragraph2_1_text = ('分析时段内'+str(typeNameList[typei])+'机型各机组自限电损失电量共'+'{:.1f}'.format(eval(metric_tongji['reason']['limTurbineLoss'])*10000)+'kWh，折合等效小时'+
                                '{:.1f}'.format(eval(metric_tongji['reason']['limTurbineLoss'])*10000/eval(windfarm_capi)/1000)+'h，各机组详细统计见下表。自限电原因分析*****')
            paragraph2_1 = document.add_paragraph()
            paragraph2_1_run = paragraph2_1.add_run(paragraph2_1_text)
            paragraph2_1_run.font.size = Pt(12)
            paragraph2_1_run.font.name = '仿宋'
            paragraph2_1_run._element.rPr.rFonts.set(qn('w:eastAsia'),u'仿宋')
            paragraph2_1.paragraph_format.first_line_indent = paragraph_run.font.size * 2
            paragraph2_1.paragraph_format.line_spacing = Pt(25)
            paragraph2_1.paragraph_format.space_after = Pt(6)
            paragraph2_1.paragraph_format.space_before = Pt(6)
            
            #表例
            table_num = table_num + 1
            paragraph2_5_text = str('表'+str(table_num)+' '+str(typeNameList[typei])+'机型自限电损失统计表')
            paragraph2_5 = document.add_paragraph()
            paragraph2_5_run = paragraph2_5.add_run(paragraph2_5_text)
            paragraph2_5_run.font.size = Pt(10)
            paragraph2_5_run.font.name = '黑体'
            paragraph2_5_run._element.rPr.rFonts.set(qn('w:eastAsia'),u'黑体')
            paragraph2_5.paragraph_format.space_after = Pt(0)
            paragraph2_5.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            #表2
            column_width = [1.9,2.7,3.5,2.7]
            device_num = len(limturbine_loss_all)+1
            tablez2_5 = document.add_table(device_num,len(column_width),style='Table Grid')
            tablez2_5.alignment = WD_TAB_ALIGNMENT.CENTER
            for i in range(device_num):
                for j in range(len(column_width)):
                    if i == 0:
                        if j == 0:
                            fill_table(tablez2_5, x=i, y=j, content='机位号', font_name='黑体', font_size=10, bold=False, paragraph_alignment=WD_PARAGRAPH_ALIGNMENT.CENTER, startendmargin=0, width=column_width[j]) 
                        elif j == 1:
                            fill_table(tablez2_5, x=i, y=j, content='限电时长(h)', font_name='黑体', font_size=10, bold=False, paragraph_alignment=WD_PARAGRAPH_ALIGNMENT.CENTER, startendmargin=0, width=column_width[j]) 
                        elif j == 2:
                            fill_table(tablez2_5, x=i, y=j, content='限电损失电量(kWh)', font_name='黑体', font_size=10, bold=False, paragraph_alignment=WD_PARAGRAPH_ALIGNMENT.CENTER, startendmargin=0, width=column_width[j]) 
                        elif j == 3:
                            fill_table(tablez2_5, x=i, y=j, content='限电时平均风速(m/s)', font_name='黑体', font_size=10, bold=False, paragraph_alignment=WD_PARAGRAPH_ALIGNMENT.CENTER, startendmargin=0, width=column_width[j]) 
                    else:
                        if j == 0:
                            fill_table(tablez2_5, x=i, y=j, content=conver_to_str(limturbine_loss_all[i-1]['wtid']), font_name='Times New Roman', font_size=10, bold=False, paragraph_alignment=WD_PARAGRAPH_ALIGNMENT.CENTER, startendmargin=0, width=column_width[j]) 
                        elif j == 1:
                            fill_table(tablez2_5, x=i, y=j, content='{:.2f}'.format(eval(limturbine_loss_all[i-1]['faultTime'])), font_name='Times New Roman', font_size=10, bold=False, paragraph_alignment=WD_PARAGRAPH_ALIGNMENT.CENTER, startendmargin=0, width=column_width[j]) 
                        elif j == 2:
                            fill_table(tablez2_5, x=i, y=j, content='{:.2f}'.format(eval(limturbine_loss_all[i-1]['faultLoss'])), font_name='Times New Roman', font_size=10, bold=False, paragraph_alignment=WD_PARAGRAPH_ALIGNMENT.CENTER, startendmargin=0, width=column_width[j]) 
                        elif j == 3:
                            fill_table(tablez2_5, x=i, y=j, content='{:.2f}'.format(eval(limturbine_loss_all[i-1]['meanWindSpeed'])), font_name='Times New Roman', font_size=10, bold=False, paragraph_alignment=WD_PARAGRAPH_ALIGNMENT.CENTER, startendmargin=0, width=column_width[j]) 
                        
        else:
            paragraph2_1_text = '分析时段内依据产业数据中台数据统计，'+str(typeNameList[typei])+'机型各机组未发生自限电情况。'      
            paragraph2_1 = document.add_paragraph()
            paragraph2_1_run = paragraph2_1.add_run(paragraph2_1_text)
            paragraph2_1_run.font.size = Pt(12)
            paragraph2_1_run.font.name = '仿宋'
            paragraph2_1_run._element.rPr.rFonts.set(qn('w:eastAsia'),u'仿宋')
            paragraph2_1.paragraph_format.first_line_indent = paragraph_run.font.size * 2
            paragraph2_1.paragraph_format.line_spacing = Pt(25)
            paragraph2_1.paragraph_format.space_after = Pt(6)
            paragraph2_1.paragraph_format.space_before = Pt(6)
    #---------------------------------------------------------------------------------------
    
    paragraph2_2_text = '3.4. 电网故障损失 '
    paragraph2_2 = document.add_paragraph()
    paragraph2_2_run = paragraph2_2.add_run(paragraph2_2_text)
    paragraph2_2_run.font.size = Pt(12)
    paragraph2_2_run.font.name = '黑体'
    paragraph2_2_run._element.rPr.rFonts.set(qn('w:eastAsia'),u'黑体')
    # paragraph2_2.paragraph_format.first_line_indent = paragraph_run.font.size * 2
    paragraph2_2.paragraph_format.space_before = Pt(12) 
    paragraph2_2.paragraph_format.space_after = Pt(0) 
    
    for typei in range(len(metric_tongjiList)): #trubine_type
        # path = str(str(path_farm)+'/'+str(np.unique(Turbine_attr['turbineTypeID'])[typei]))
        metric_tongji = metric_tongjiList[typei]#pd.read_csv(str(path+'/tongji.csv'),header=[0],index_col=[0])
        
        # path = str(str(path_farm)+'/'+str(np.unique(Turbine_attr['turbineTypeID'])[typei]))
        faultgrid_loss_all = show_grid_fault_loss.analyse(farmInfo['farm_name'], [typeNameList[typei]], startTime, endTime)#['table']#pd.read_csv(str(path+'/faultgrid_loss_all.csv'),header=[0],index_col=[0])
        if len(faultgrid_loss_all) > 0:
            faultgrid_loss_all = faultgrid_loss_all['table']
        
        if (eval(metric_tongji['reason']['gridFaultLoss'])*10000 > 200) and (len(faultgrid_loss_all)>0):
            paragraph2_1_text = ('分析时段内'+str(typeNameList[typei])+'机型各机组因电网故障发生的损失电量共'+'{:.1f}'.format(eval(metric_tongji['reason']['gridFaultLoss'])*10000)+'kWh，折合等效小时'+
                                '{:.1f}'.format(eval(metric_tongji['reason']['gridFaultLoss'])*10000/eval(windfarm_capi)/1000)+'h，详细故障损失统计见下表，其中***故障发生频次较高。')   
            paragraph2_1 = document.add_paragraph()
            paragraph2_1_run = paragraph2_1.add_run(paragraph2_1_text)
            paragraph2_1_run.font.size = Pt(12)
            paragraph2_1_run.font.name = '仿宋'
            paragraph2_1_run._element.rPr.rFonts.set(qn('w:eastAsia'),u'仿宋')
            paragraph2_1.paragraph_format.first_line_indent = paragraph_run.font.size * 2
            paragraph2_1.paragraph_format.line_spacing = Pt(25)
            paragraph2_1.paragraph_format.space_after = Pt(6)
            paragraph2_1.paragraph_format.space_before = Pt(6)
            
            #表例
            table_num = table_num + 1
            paragraph2_5_text = str('表'+str(table_num)+' '+str(typeNameList[typei])+'电网故障损失统计表')
            paragraph2_5 = document.add_paragraph()
            paragraph2_5_run = paragraph2_5.add_run(paragraph2_5_text)
            paragraph2_5_run.font.size = Pt(10)
            paragraph2_5_run.font.name = '黑体'
            paragraph2_5_run._element.rPr.rFonts.set(qn('w:eastAsia'),u'黑体')
            paragraph2_5.paragraph_format.space_after = Pt(0)
            paragraph2_5.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            #表2
            column_width = [1.25,1.70,1.72,2.35,3.5,3.99]
            device_num = len(faultgrid_loss_all)+1
            tablez2_5 = document.add_table(device_num,len(column_width),style='Table Grid')
            tablez2_5.alignment = WD_TAB_ALIGNMENT.CENTER
            for i in range(device_num):
                for j in range(len(column_width)):
                    if i == 0:
                        if j == 0:
                            fill_table(tablez2_5, x=i, y=j, content='机位号', font_name='黑体', font_size=10, bold=False, paragraph_alignment=WD_PARAGRAPH_ALIGNMENT.CENTER, startendmargin=0, width=column_width[j]) 
                        elif j == 1:
                            fill_table(tablez2_5, x=i, y=j, content='故障码', font_name='黑体', font_size=10, bold=False, paragraph_alignment=WD_PARAGRAPH_ALIGNMENT.CENTER, startendmargin=0, width=column_width[j]) 
                        elif j == 2:
                            fill_table(tablez2_5, x=i, y=j, content='故障频次', font_name='黑体', font_size=10, bold=False, paragraph_alignment=WD_PARAGRAPH_ALIGNMENT.CENTER, startendmargin=0, width=column_width[j]) 
                        elif j == 3:
                            fill_table(tablez2_5, x=i, y=j, content='故障时长(h)', font_name='黑体', font_size=10, bold=False, paragraph_alignment=WD_PARAGRAPH_ALIGNMENT.CENTER, startendmargin=0, width=column_width[j]) 
                        elif j == 4:
                            fill_table(tablez2_5, x=i, y=j, content='故障损失电量(kWh)', font_name='黑体', font_size=10, bold=False, paragraph_alignment=WD_PARAGRAPH_ALIGNMENT.CENTER, startendmargin=0, width=column_width[j]) 
                        elif j == 5:
                            fill_table(tablez2_5, x=i, y=j, content='故障描述', font_name='黑体', font_size=10, bold=False, paragraph_alignment=WD_PARAGRAPH_ALIGNMENT.CENTER, startendmargin=0, width=column_width[j]) 
                    else:
                        if j == 0:
                            fill_table(tablez2_5, x=i, y=j, content=conver_to_str(faultgrid_loss_all[i-1]['wtid']), font_name='Times New Roman', font_size=10, bold=False, paragraph_alignment=WD_PARAGRAPH_ALIGNMENT.CENTER, startendmargin=0, width=column_width[j]) 
                        elif j == 1:
                            fill_table(tablez2_5, x=i, y=j, content=int(faultgrid_loss_all[i-1]['faultCode']), font_name='Times New Roman', font_size=10, bold=False, paragraph_alignment=WD_PARAGRAPH_ALIGNMENT.CENTER, startendmargin=0, width=column_width[j]) 
                        elif j == 2:
                            fill_table(tablez2_5, x=i, y=j, content=int(faultgrid_loss_all[i-1]['faultCount']), font_name='Times New Roman', font_size=10, bold=False, paragraph_alignment=WD_PARAGRAPH_ALIGNMENT.CENTER, startendmargin=0, width=column_width[j]) 
                        elif j == 3:
                            fill_table(tablez2_5, x=i, y=j, content='{:.2f}'.format(eval(faultgrid_loss_all[i-1]['faultTime'])), font_name='Times New Roman', font_size=10, bold=False, paragraph_alignment=WD_PARAGRAPH_ALIGNMENT.CENTER, startendmargin=0, width=column_width[j]) 
                        elif j == 4:
                            fill_table(tablez2_5, x=i, y=j, content='{:.2f}'.format(eval(faultgrid_loss_all[i-1]['faultLoss'])), font_name='Times New Roman', font_size=10, bold=False, paragraph_alignment=WD_PARAGRAPH_ALIGNMENT.CENTER, startendmargin=0, width=column_width[j]) 
                        elif j == 5:
                            fill_table(tablez2_5, x=i, y=j, content=str(faultgrid_loss_all[i-1]['faultDescribe']), font_name='仿宋', font_size=10, bold=False, paragraph_alignment=WD_PARAGRAPH_ALIGNMENT.CENTER, startendmargin=0, width=column_width[j]) 
        else:
            paragraph2_1_text = '分析时段内'+str(typeNameList[typei])+'机型各机组未发生电网侧故障。'      
            paragraph2_1 = document.add_paragraph()
            paragraph2_1_run = paragraph2_1.add_run(paragraph2_1_text)
            paragraph2_1_run.font.size = Pt(12)
            paragraph2_1_run.font.name = '仿宋'
            paragraph2_1_run._element.rPr.rFonts.set(qn('w:eastAsia'),u'仿宋')
            paragraph2_1.paragraph_format.first_line_indent = paragraph_run.font.size * 2
            paragraph2_1.paragraph_format.line_spacing = Pt(25)
            paragraph2_1.paragraph_format.space_after = Pt(6)
            paragraph2_1.paragraph_format.space_before = Pt(6)


            
    paragraph2_2_text = '3.5. 技术待机损失 '
    paragraph2_2 = document.add_paragraph()
    paragraph2_2_run = paragraph2_2.add_run(paragraph2_2_text)
    paragraph2_2_run.font.size = Pt(12)
    paragraph2_2_run.font.name = '黑体'
    paragraph2_2_run._element.rPr.rFonts.set(qn('w:eastAsia'),u'黑体')
    # paragraph2_2.paragraph_format.first_line_indent = paragraph_run.font.size * 2
    paragraph2_2.paragraph_format.space_before = Pt(12) 
    paragraph2_2.paragraph_format.space_after = Pt(0) 
    
    paragraph2_1_text = ('技术待机是由于机组自身运行、控制机制造成的待机。')   
    paragraph2_1 = document.add_paragraph()
    paragraph2_1_run = paragraph2_1.add_run(paragraph2_1_text)
    paragraph2_1_run.font.size = Pt(12)
    paragraph2_1_run.font.name = '仿宋'
    paragraph2_1_run._element.rPr.rFonts.set(qn('w:eastAsia'),u'仿宋')
    paragraph2_1.paragraph_format.first_line_indent = paragraph_run.font.size * 2
    paragraph2_1.paragraph_format.line_spacing = Pt(25)
    paragraph2_1.paragraph_format.space_after = Pt(0)
    paragraph2_1.paragraph_format.space_before = Pt(6)
    
    for typei in range(len(metric_tongjiList)): # trubine_type
        # path = str(str(path_farm)+'/'+str(np.unique(Turbine_attr['turbineTypeID'])[typei]))
        metric_tongji = metric_tongjiList[typei] #pd.read_csv(str(path+'/tongji.csv'),header=[0],index_col=[0])
        
        # path = str(str(path_farm)+'/'+str(np.unique(Turbine_attr['turbineTypeID'])[typei]))
        Technology_loss_all = show_technology_loss.analyse(farmInfo['farm_name'], [typeNameList[typei]], startTime, endTime)#['table'] #pd.read_csv(str(path+'/Technology_loss_all.csv'),header=[0],index_col=[0])
        if len(Technology_loss_all) > 0:
            Technology_loss_all = Technology_loss_all['table']
        
        if (eval(metric_tongji['reason']['technologyLoss'])*10000 > 200) and (len(Technology_loss_all)>0):
            paragraph2_1_text = ('分析时段内'+str(typeNameList[typei])+'机型各机组技术待机损失电量共'+'{:.1f}'.format(eval(metric_tongji.loc[0]['Technology_loss']))+'kWh，折合等效小时'+
                                '{:.1f}'.format(eval(metric_tongji['reason']['technologyLoss'])*10000/eval(windfarm_capi)/1000)+'h，详细损失统计见下表。')   
            paragraph2_1 = document.add_paragraph()
            paragraph2_1_run = paragraph2_1.add_run(paragraph2_1_text)
            paragraph2_1_run.font.size = Pt(12)
            paragraph2_1_run.font.name = '仿宋'
            paragraph2_1_run._element.rPr.rFonts.set(qn('w:eastAsia'),u'仿宋')
            paragraph2_1.paragraph_format.first_line_indent = paragraph_run.font.size * 2
            paragraph2_1.paragraph_format.line_spacing = Pt(25)
            paragraph2_1.paragraph_format.space_after = Pt(6)
            paragraph2_1.paragraph_format.space_before = Pt(0)
            
            #表例
            table_num = table_num + 1
            paragraph2_5_text = str('表'+str(table_num)+' '+str(typeNameList[typei])+'机型技术待机损失统计表')
            paragraph2_5 = document.add_paragraph()
            paragraph2_5_run = paragraph2_5.add_run(paragraph2_5_text)
            paragraph2_5_run.font.size = Pt(10)
            paragraph2_5_run.font.name = '黑体'
            paragraph2_5_run._element.rPr.rFonts.set(qn('w:eastAsia'),u'黑体')
            paragraph2_5.paragraph_format.space_after = Pt(0)
            paragraph2_5.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            #表2
            column_width = [1.25,1.70,1.85,3.5,2.67,3.99]
            device_num = len(Technology_loss_all)+1
            tablez2_5 = document.add_table(device_num,len(column_width),style='Table Grid')
            tablez2_5.alignment = WD_TAB_ALIGNMENT.CENTER
            for i in range(device_num):
                for j in range(len(column_width)):
                    if i == 0:
                        if j == 0:
                            fill_table(tablez2_5, x=i, y=j, content='机位号', font_name='黑体', font_size=10, bold=False, paragraph_alignment=WD_PARAGRAPH_ALIGNMENT.CENTER, startendmargin=0, width=column_width[j]) 
                        elif j == 1:
                            fill_table(tablez2_5, x=i, y=j, content='待机频次', font_name='黑体', font_size=10, bold=False, paragraph_alignment=WD_PARAGRAPH_ALIGNMENT.CENTER, startendmargin=0, width=column_width[j]) 
                        elif j == 2:
                            fill_table(tablez2_5, x=i, y=j, content='待机时长(h)', font_name='黑体', font_size=10, bold=False, paragraph_alignment=WD_PARAGRAPH_ALIGNMENT.CENTER, startendmargin=0, width=column_width[j]) 
                        elif j == 3:
                            fill_table(tablez2_5, x=i, y=j, content='待机损失电量(kWh)', font_name='黑体', font_size=10, bold=False, paragraph_alignment=WD_PARAGRAPH_ALIGNMENT.CENTER, startendmargin=0, width=column_width[j]) 
                        elif j == 4:
                            fill_table(tablez2_5, x=i, y=j, content='待机时平均风速(m/s)', font_name='黑体', font_size=10, bold=False, paragraph_alignment=WD_PARAGRAPH_ALIGNMENT.CENTER, startendmargin=0, width=column_width[j]) 
                        elif j == 5:
                            fill_table(tablez2_5, x=i, y=j, content='待机原因', font_name='黑体', font_size=10, bold=False, paragraph_alignment=WD_PARAGRAPH_ALIGNMENT.CENTER, startendmargin=0, width=column_width[j]) 
                    else:
                        if j == 0:
                            fill_table(tablez2_5, x=i, y=j, content=conver_to_str(Technology_loss_all[i-1]['wtid']), font_name='Times New Roman', font_size=10, bold=False, paragraph_alignment=WD_PARAGRAPH_ALIGNMENT.CENTER, startendmargin=0, width=column_width[j]) 
                        elif j == 1:
                            fill_table(tablez2_5, x=i, y=j, content=int(Technology_loss_all[i-1]['faultCount']), font_name='Times New Roman', font_size=10, bold=False, paragraph_alignment=WD_PARAGRAPH_ALIGNMENT.CENTER, startendmargin=0, width=column_width[j]) 
                        elif j == 2:
                            fill_table(tablez2_5, x=i, y=j, content='{:.2f}'.format(eval(Technology_loss_all[i-1]['faultTime'])), font_name='Times New Roman', font_size=10, bold=False, paragraph_alignment=WD_PARAGRAPH_ALIGNMENT.CENTER, startendmargin=0, width=column_width[j]) 
                        elif j == 3:
                            fill_table(tablez2_5, x=i, y=j, content='{:.2f}'.format(eval(Technology_loss_all[i-1]['faultLoss'])), font_name='Times New Roman', font_size=10, bold=False, paragraph_alignment=WD_PARAGRAPH_ALIGNMENT.CENTER, startendmargin=0, width=column_width[j]) 
                        elif j == 4:
                            fill_table(tablez2_5, x=i, y=j, content='{:.2f}'.format(eval(Technology_loss_all[i-1]['meanWindSpeed'])), font_name='Times New Roman', font_size=10, bold=False, paragraph_alignment=WD_PARAGRAPH_ALIGNMENT.CENTER, startendmargin=0, width=column_width[j]) 
                        elif j == 5:
                            fill_table(tablez2_5, x=i, y=j, content=str(Technology_loss_all[i-1]['faultDescribe']), font_name='仿宋', font_size=10, bold=False, paragraph_alignment=WD_PARAGRAPH_ALIGNMENT.CENTER, startendmargin=0, width=column_width[j]) 
        else:
            paragraph2_1_text = '分析时段内'+str(typeNameList[typei])+'机型各机组未发生技术待机情况。'      
            paragraph2_1 = document.add_paragraph()
            paragraph2_1_run = paragraph2_1.add_run(paragraph2_1_text)
            paragraph2_1_run.font.size = Pt(12)
            paragraph2_1_run.font.name = '仿宋'
            paragraph2_1_run._element.rPr.rFonts.set(qn('w:eastAsia'),u'仿宋')
            paragraph2_1.paragraph_format.first_line_indent = paragraph_run.font.size * 2
            paragraph2_1.paragraph_format.line_spacing = Pt(25)
            paragraph2_1.paragraph_format.space_after = Pt(6)
            paragraph2_1.paragraph_format.space_before = Pt(6)
    #---------------------------------------------------------------------------------------
    
    paragraph2_2_text = '3.6. 计划停机损失 '
    paragraph2_2 = document.add_paragraph()
    paragraph2_2_run = paragraph2_2.add_run(paragraph2_2_text)
    paragraph2_2_run.font.size = Pt(12)
    paragraph2_2_run.font.name = '黑体'
    paragraph2_2_run._element.rPr.rFonts.set(qn('w:eastAsia'),u'黑体')
    # paragraph2_2.paragraph_format.first_line_indent = paragraph_run.font.size * 2
    paragraph2_2.paragraph_format.space_before = Pt(12) 
    paragraph2_2.paragraph_format.space_after = Pt(0) 
    
    paragraph2_1_text = ('计划停机损失是机组手动停机、维护状态停机或机组无故障停机且风速大于并网发电风速时导致的发电量损失。')   
    paragraph2_1 = document.add_paragraph()
    paragraph2_1_run = paragraph2_1.add_run(paragraph2_1_text)
    paragraph2_1_run.font.size = Pt(12)
    paragraph2_1_run.font.name = '仿宋'
    paragraph2_1_run._element.rPr.rFonts.set(qn('w:eastAsia'),u'仿宋')
    paragraph2_1.paragraph_format.first_line_indent = paragraph_run.font.size * 2
    paragraph2_1.paragraph_format.line_spacing = Pt(25)
    paragraph2_1.paragraph_format.space_after = Pt(0)
    paragraph2_1.paragraph_format.space_before = Pt(6)
    
    for typei in range(len(metric_tongjiList)): # trubine_type
        # path = str(str(path_farm)+'/'+str(np.unique(Turbine_attr['turbineTypeID'])[typei]))
        metric_tongji = metric_tongjiList[typei]#pd.read_csv(str(path+'/tongji.csv'),header=[0],index_col=[0])
        
        # path = str(str(path_farm)+'/'+str(np.unique(Turbine_attr['turbineTypeID'])[typei]))
        stop_loss_all = show_stop_loss.analyse(farmInfo['farm_name'], [typeNameList[typei]], startTime, endTime)#['table']#pd.read_csv(str(path+'/stop_loss_all.csv'),header=[0],index_col=[0])
        if len(stop_loss_all) > 0:
            stop_loss_all = stop_loss_all['table']
        
        if (eval(metric_tongji['reason']['stopLoss'])*10000 > 200) and (len(stop_loss_all)>0):
            paragraph2_1_text = ('分析时段内'+str(typeNameList[typei])+'机型各机组计划停机损失电量共'+'{:.1f}'.format(eval(metric_tongji.loc[0]['stop_loss']))+'kWh，折合等效小时'+
                                '{:.1f}'.format(eval(metric_tongji['reason']['stopLoss'])*10000/eval(windfarm_capi)/1000)+'h，详细损失统计见下图表。')   
            paragraph2_1 = document.add_paragraph()
            paragraph2_1_run = paragraph2_1.add_run(paragraph2_1_text)
            paragraph2_1_run.font.size = Pt(12)
            paragraph2_1_run.font.name = '仿宋'
            paragraph2_1_run._element.rPr.rFonts.set(qn('w:eastAsia'),u'仿宋')
            paragraph2_1.paragraph_format.first_line_indent = paragraph_run.font.size * 2
            paragraph2_1.paragraph_format.line_spacing = Pt(25)
            paragraph2_1.paragraph_format.space_after = Pt(6)
            paragraph2_1.paragraph_format.space_before = Pt(0)
            
            #表例
            table_num = table_num + 1
            paragraph2_5_text = str('表'+str(table_num)+' '+str(typeNameList[typei])+'机型计划停机损失统计表')
            paragraph2_5 = document.add_paragraph()
            paragraph2_5_run = paragraph2_5.add_run(paragraph2_5_text)
            paragraph2_5_run.font.size = Pt(10)
            paragraph2_5_run.font.name = '黑体'
            paragraph2_5_run._element.rPr.rFonts.set(qn('w:eastAsia'),u'黑体')
            paragraph2_5.paragraph_format.space_after = Pt(0)
            paragraph2_5.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            #表2
            column_width = [1.25,1.70,3.5,2.67,2.7]
            device_num = len(stop_loss_all)+1
            tablez2_5 = document.add_table(device_num,len(column_width),style='Table Grid')
            tablez2_5.alignment = WD_TAB_ALIGNMENT.CENTER
            for i in range(device_num):
                for j in range(len(column_width)):
                    if i == 0:
                        if j == 0:
                            fill_table(tablez2_5, x=i, y=j, content='机位号', font_name='黑体', font_size=10, bold=False, paragraph_alignment=WD_PARAGRAPH_ALIGNMENT.CENTER, startendmargin=0, width=column_width[j]) 
                        elif j == 1:
                            fill_table(tablez2_5, x=i, y=j, content='停机时长(h)', font_name='黑体', font_size=10, bold=False, paragraph_alignment=WD_PARAGRAPH_ALIGNMENT.CENTER, startendmargin=0, width=column_width[j]) 
                        elif j == 2:
                            fill_table(tablez2_5, x=i, y=j, content='停机损失电量(kWh)', font_name='黑体', font_size=10, bold=False, paragraph_alignment=WD_PARAGRAPH_ALIGNMENT.CENTER, startendmargin=0, width=column_width[j]) 
                        elif j == 3:
                            fill_table(tablez2_5, x=i, y=j, content='停机时平均风速(m/s)', font_name='黑体', font_size=10, bold=False, paragraph_alignment=WD_PARAGRAPH_ALIGNMENT.CENTER, startendmargin=0, width=column_width[j]) 
                        elif j == 4:
                            fill_table(tablez2_5, x=i, y=j, content='停机时环境温度(℃)', font_name='黑体', font_size=10, bold=False, paragraph_alignment=WD_PARAGRAPH_ALIGNMENT.CENTER, startendmargin=0, width=column_width[j]) 
                    else:
                        if j == 0:
                            fill_table(tablez2_5, x=i, y=j, content=conver_to_str(stop_loss_all[i-1]['wtid']), font_name='Times New Roman', font_size=10, bold=False, paragraph_alignment=WD_PARAGRAPH_ALIGNMENT.CENTER, startendmargin=0, width=column_width[j]) 
                        elif j == 1:
                            fill_table(tablez2_5, x=i, y=j, content='{:.2f}'.format(eval(stop_loss_all.iloc[i-1]['faultTime'])), font_name='Times New Roman', font_size=10, bold=False, paragraph_alignment=WD_PARAGRAPH_ALIGNMENT.CENTER, startendmargin=0, width=column_width[j]) 
                        elif j == 2:
                            fill_table(tablez2_5, x=i, y=j, content='{:.2f}'.format(eval(stop_loss_all.iloc[i-1]['faultLoss'])), font_name='Times New Roman', font_size=10, bold=False, paragraph_alignment=WD_PARAGRAPH_ALIGNMENT.CENTER, startendmargin=0, width=column_width[j]) 
                        elif j == 3:
                            fill_table(tablez2_5, x=i, y=j, content='{:.2f}'.format(eval(stop_loss_all.iloc[i-1]['meanWindSpeed'])), font_name='Times New Roman', font_size=10, bold=False, paragraph_alignment=WD_PARAGRAPH_ALIGNMENT.CENTER, startendmargin=0, width=column_width[j]) 
                        elif j == 4:
                            fill_table(tablez2_5, x=i, y=j, content='{:.2f}'.format(eval(stop_loss_all[i-1]['exltmp'])), font_name='Times New Roman', font_size=10, bold=False, paragraph_alignment=WD_PARAGRAPH_ALIGNMENT.CENTER, startendmargin=0, width=column_width[j]) 
        else:
            paragraph2_1_text = '分析时段内'+str(typeNameList[typei])+'机型各机组未发生计划停机情况。'      
            paragraph2_1 = document.add_paragraph()
            paragraph2_1_run = paragraph2_1.add_run(paragraph2_1_text)
            paragraph2_1_run.font.size = Pt(12)
            paragraph2_1_run.font.name = '仿宋'
            paragraph2_1_run._element.rPr.rFonts.set(qn('w:eastAsia'),u'仿宋')
            paragraph2_1.paragraph_format.first_line_indent = paragraph_run.font.size * 2
            paragraph2_1.paragraph_format.line_spacing = Pt(25)
            paragraph2_1.paragraph_format.space_after = Pt(6)
            paragraph2_1.paragraph_format.space_before = Pt(6)
#---------------------------------------------------------------------------------------
    #小节4
    paragraph2 = document.add_paragraph(style='ListNumber')
    paragraph2_run = paragraph2.add_run('功率曲线')
    paragraph2_run.font.size = Pt(14)
    paragraph2_run.font.name = '黑体'
    paragraph2_run._element.rPr.rFonts.set(qn('w:eastAsia'),u'黑体')
    paragraph2.paragraph_format.space_before = Pt(12)
    paragraph2.paragraph_format.space_after = Pt(6)
    

    #文字
    for i in range(len(metric_tongjiList)): # trubine_type
        # path = str(str(path_farm)+'/'+str(np.unique(Turbine_attr['turbineTypeID'])[i]))
        try:
            result, power_tongji = show_power_consistence.analyse(farmInfo['farm_name'], typeNameList[i], farmInfo['wtid'][typeNameList[i]], startTime, endTime)#pd.read_csv(str(path+'/turbine_err_all.csv'),header=[0],index_col=[0])
        except Exception as e:
            errorInfomation = traceback.format_exc()
            # print("数据库中表eny_wspd_all在时段"+str(startTime)+"到"+str(endTime)+"没有机型"+str(typeNameList[i-1]))
            print("############################数据库中表pw_turbine_all在时段"+str(startTime)+"到"+str(endTime)+"没有机型"+str(typeNameList[i])+'的数据导致报错#################################')
            print(f'{errorInfomation}')
            continue
        
        if len(power_tongji) <= 5:
            top_num = 1
        elif (len(power_tongji) > 5) and (len(power_tongji) <= 10):
            top_num = 2
        elif (len(power_tongji) > 10) and (len(power_tongji) <= 20):
            top_num = 3
        else: 
            top_num = 5
            
        wtid_top = power_tongji.loc[power_tongji['k_order'].nlargest(top_num).index,'wtid']
        wtid_low = power_tongji.loc[power_tongji['k_order'].nsmallest(top_num).index,'wtid']
        
        if (np.std(power_tongji['k_order'])>0.05)|(len(power_tongji[(power_tongji['k_order']>1.05)|(power_tongji['k_order']<0.95)])>0):
            yizhi = '较差'
        else:
            yizhi = '较好'
    
        paragraph2_1_text = ('分析时段内，'+windfarm_name+str(typeNameList[i])+'机型功率曲线一致性'+yizhi+
                             str('，其中功率曲线表现较好的机组为')+str('、'.join(wtid_top.astype(str)))+
                             str('，功率曲线表现较差的机组为')+str('、'.join(wtid_low.astype(str)))+'。')

        paragraph2_1 = document.add_paragraph()
        paragraph2_1_run = paragraph2_1.add_run(paragraph2_1_text)
        paragraph2_1_run.font.size = Pt(12)
        paragraph2_1_run.font.name = '仿宋'
        paragraph2_1_run._element.rPr.rFonts.set(qn('w:eastAsia'),u'仿宋')
        paragraph2_1.paragraph_format.first_line_indent = paragraph_run.font.size * 2
        paragraph2_1.paragraph_format.line_spacing = Pt(25)
        paragraph2_1.paragraph_format.space_after = Pt(6)
        
        #图补
        wtids = farmInfo['wtid'][typeNameList[i]]#Turbine_attr[Turbine_attr['turbineTypeID']==np.unique(Turbine_attr['turbineTypeID'])[i]]
        turbine_num = 16
        pnum = len(wtids)//turbine_num
        rem = len(wtids)%turbine_num
        print('pnum'+'//'+str(pnum))
        if pnum > 0: 
            if rem > 0:
                pnum_new = pnum + 1
            else:
                pnum_new = pnum
            #rem_pw = turbine_num
            for j_pw in range(pnum_new):
                print(j_pw)
                #fig.savefig(path + '/' +'功率曲线'+str(j_pw) +'.png',dpi=100, transparent=True, bbox_inches='tight')
                figure_path = selectPowerCurvePicture(farmInfo['farm_name'], typeNameList[i], '功率曲线'+str(j_pw) +'.png', startTime, endTime)#str(str(path)+'/'+'/'+'功率曲线'+str(j_pw) +'.png')
                if figure_path == None:
                    continue
                paragraph_picture2_4 = document.add_paragraph()
                paragraph_picture2_4.alignment = WD_ALIGN_PARAGRAPH.CENTER
                paragraph_picture2_4_run = paragraph_picture2_4.add_run()
                paragraph_picture2_4_run.add_picture(figure_path,width=Cm(14))
                paragraph_picture2_4.paragraph_format.space_after = Pt(0)
                #图示
                picture_num += 1
                paragraph2_4_text = str('图'+str(picture_num)+' '+str(typeNameList[i])+'机型功率曲线'+str(j_pw+1))
                paragraph2_4 = document.add_paragraph()
                paragraph2_4_run = paragraph2_4.add_run(paragraph2_4_text)
                paragraph2_4_run.font.size = Pt(10)
                paragraph2_4_run.font.name = '黑体'
                paragraph2_4_run._element.rPr.rFonts.set(qn('w:eastAsia'),u'黑体')
                paragraph2_4.paragraph_format.space_after = Pt(16)
                paragraph2_4.alignment = WD_ALIGN_PARAGRAPH.CENTER
                
                figure_path = selectCPPicture(farmInfo['farm_name'], typeNameList[i], 'Cp曲线'+str(j_pw) +'.png', startTime, endTime) #str(str(path)+'/'+'/'+'Cp曲线'+str(j_pw) +'.png')
                if figure_path == None:
                    continue
                paragraph_picture2_4 = document.add_paragraph()
                paragraph_picture2_4.alignment = WD_ALIGN_PARAGRAPH.CENTER
                paragraph_picture2_4_run = paragraph_picture2_4.add_run()
                paragraph_picture2_4_run.add_picture(figure_path,width=Cm(14))
                paragraph_picture2_4.paragraph_format.space_after = Pt(0)
                #图示
                picture_num += 1
                paragraph2_4_text = str('图'+str(picture_num)+' '+str(typeNameList[i])+'机型Cp曲线'+str(j_pw+1))
                paragraph2_4 = document.add_paragraph()
                paragraph2_4_run = paragraph2_4.add_run(paragraph2_4_text)
                paragraph2_4_run.font.size = Pt(10)
                paragraph2_4_run.font.name = '黑体'
                paragraph2_4_run._element.rPr.rFonts.set(qn('w:eastAsia'),u'黑体')
                paragraph2_4.paragraph_format.space_after = Pt(16)
                paragraph2_4.alignment = WD_ALIGN_PARAGRAPH.CENTER
        else:
            #fig.savefig(path + '/' +'功率曲线'+str(pnum+1) +'.png',dpi=100,transparent=True, bbox_inches='tight')
            figure_path = selectPowerCurvePicture(farmInfo['farm_name'], typeNameList[i], '功率曲线'+str(pnum+1) +'.png', startTime, endTime)#str(str(path)+'/'+'/'+'功率曲线'+str(pnum+1) +'.png')
            if figure_path == None:
                    continue
            paragraph_picture2_4 = document.add_paragraph()
            paragraph_picture2_4.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph_picture2_4_run = paragraph_picture2_4.add_run()
            paragraph_picture2_4_run.add_picture(figure_path,width=Cm(14))
            paragraph_picture2_4.paragraph_format.space_after = Pt(0)
            #图示
            picture_num += 1
            paragraph2_4_text = str('图'+str(picture_num)+' '+str(typeNameList[i])+'机型功率曲线')
            paragraph2_4 = document.add_paragraph()
            paragraph2_4_run = paragraph2_4.add_run(paragraph2_4_text)
            paragraph2_4_run.font.size = Pt(10)
            paragraph2_4_run.font.name = '黑体'
            paragraph2_4_run._element.rPr.rFonts.set(qn('w:eastAsia'),u'黑体')
            paragraph2_4.paragraph_format.space_after = Pt(16)
            paragraph2_4.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            figure_path = selectCPPicture(farmInfo['farm_name'], typeNameList[i], 'Cp曲线'+str(pnum+1) +'.png', startTime, endTime) #str(str(path)+'/'+'/'+'Cp曲线'+str(pnum+1) +'.png')
            if figure_path == None:
                    continue
            paragraph_picture2_4 = document.add_paragraph()
            paragraph_picture2_4.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph_picture2_4_run = paragraph_picture2_4.add_run()
            paragraph_picture2_4_run.add_picture(figure_path,width=Cm(14))
            paragraph_picture2_4.paragraph_format.space_after = Pt(0)
            #图示
            picture_num += 1
            paragraph2_4_text = str('图'+str(picture_num)+' '+str(typeNameList[i])+'机型Cp曲线')
            paragraph2_4 = document.add_paragraph()
            paragraph2_4_run = paragraph2_4.add_run(paragraph2_4_text)
            paragraph2_4_run.font.size = Pt(10)
            paragraph2_4_run.font.name = '黑体'
            paragraph2_4_run._element.rPr.rFonts.set(qn('w:eastAsia'),u'黑体')
            paragraph2_4.paragraph_format.space_after = Pt(16)
            paragraph2_4.alignment = WD_ALIGN_PARAGRAPH.CENTER


#---------------------------------------------------------------------------------------
    #小节4
    paragraph2 = document.add_paragraph(style='ListNumber')
    paragraph2_run = paragraph2.add_run('控制性能分析')
    paragraph2_run.font.size = Pt(14)
    paragraph2_run.font.name = '黑体'
    paragraph2_run._element.rPr.rFonts.set(qn('w:eastAsia'),u'黑体')
    paragraph2.paragraph_format.space_before = Pt(12)
    paragraph2.paragraph_format.space_after = Pt(6) 
    
    #文字
    paragraph1_5_text = '从产业数据中台提取了各机组的转速、桨距角、功率、风速、风向等测点数据，从偏航控制、变桨控制、转矩控制等方面分析控制策略对机组发电性能的影响。'
    paragraph1_5 = document.add_paragraph()
    paragraph1_5_run = paragraph1_5.add_run(paragraph1_5_text)
    paragraph1_5_run.font.size = Pt(12)
    paragraph1_5_run.font.name = '仿宋'
    paragraph1_5_run._element.rPr.rFonts.set(qn('w:eastAsia'),u'仿宋')
    paragraph1_5.paragraph_format.first_line_indent = paragraph_run.font.size * 2
    paragraph1_5.paragraph_format.line_spacing = Pt(25)
    paragraph1_5.paragraph_format.space_after = Pt(6)


    paragraph2_2_text = '5.1. 偏航控制分析 '
    paragraph2_2 = document.add_paragraph()
    paragraph2_2_run = paragraph2_2.add_run(paragraph2_2_text)
    paragraph2_2_run.font.size = Pt(12)
    paragraph2_2_run.font.name = '黑体'
    paragraph2_2_run._element.rPr.rFonts.set(qn('w:eastAsia'),u'黑体')
    # paragraph2_2.paragraph_format.first_line_indent = paragraph_run.font.size * 2
    paragraph2_2.paragraph_format.space_before = Pt(12) 
    paragraph2_2.paragraph_format.space_after = Pt(0) 
    
    paragraph2_1_text = ('偏航控制常见问题为偏航对风不正，产生偏航对风不正的原因通常为：1)风向标标定不准；2)风向标安装松动；3)叶轮尾流影响；4)偏航控制策略。当偏航误差角度超过5度，将对机组发电量造成明显影响。')   
    paragraph2_1 = document.add_paragraph()
    paragraph2_1_run = paragraph2_1.add_run(paragraph2_1_text)
    paragraph2_1_run.font.size = Pt(12)
    paragraph2_1_run.font.name = '仿宋'
    paragraph2_1_run._element.rPr.rFonts.set(qn('w:eastAsia'),u'仿宋')
    paragraph2_1.paragraph_format.first_line_indent = paragraph_run.font.size * 2
    paragraph2_1.paragraph_format.line_spacing = Pt(25)
    paragraph2_1.paragraph_format.space_after = Pt(0)
    paragraph2_1.paragraph_format.space_before = Pt(6)
    
    for typei in range(len(metric_tongjiList)):#trubine_type
        
        # path = str(str(path_farm)+'/'+str(np.unique(Turbine_attr['turbineTypeID'])[typei]))
        err_result_great5 = []#pd.read_csv(str(path+'/err_result_all.csv'),header=[0],index_col=[0])
        err_result_less5 = []
        # print(path)
        for turbineName in farmInfo['wtid'][typeNameList[typei]]:
            resultBiasDirection = selectNavigationBiasDirectionPicture(farmInfo['farm_name'], typeNameList[typei], turbineName, startTime, endTime)
            if resultBiasDirection is not None:
                dictBiasDirection = {}
                dictBiasDirection['turbine'] = resultBiasDirection[0]
                dictBiasDirection['yawerr'] = resultBiasDirection[2]
                dictBiasDirection['picture'] = resultBiasDirection[1]
                dictBiasDirection['loss'] = resultBiasDirection[3]
                if np.abs(dictBiasDirection['yawerr']) > 5.0: # and np.abs(dictBiasDirection['yawerr']) < 999
                    err_result_great5.append(dictBiasDirection)
                else:
                    err_result_less5.append(dictBiasDirection)

        if len(err_result_great5) > 0:
            # err_temp = err_result_all[(np.abs(err_result_all['yawerr'])>5.0)&(np.abs(err_result_all['yawerr'])<999)]
            #err_temp = err_temp.reset_index()
            # if len(err_temp)>0:
            paragraph2_1_text = ('分析时段内'+str(typeNameList[typei])+'机型部分机组存在偏航对风不正问题，详细见下表。')   
            paragraph2_1 = document.add_paragraph()
            paragraph2_1_run = paragraph2_1.add_run(paragraph2_1_text)
            paragraph2_1_run.font.size = Pt(12)
            paragraph2_1_run.font.name = '仿宋'
            paragraph2_1_run._element.rPr.rFonts.set(qn('w:eastAsia'),u'仿宋')
            paragraph2_1.paragraph_format.first_line_indent = paragraph_run.font.size * 2
            paragraph2_1.paragraph_format.line_spacing = Pt(25)
            paragraph2_1.paragraph_format.space_after = Pt(6)
            paragraph2_1.paragraph_format.space_before = Pt(6)
            
            #表例
            table_num = table_num + 1
            paragraph2_5_text = str('表'+str(table_num)+' '+str(typeNameList[typei])+'机型偏航对风不正机组统计表')
            paragraph2_5 = document.add_paragraph()
            paragraph2_5_run = paragraph2_5.add_run(paragraph2_5_text)
            paragraph2_5_run.font.size = Pt(10)
            paragraph2_5_run.font.name = '黑体'
            paragraph2_5_run._element.rPr.rFonts.set(qn('w:eastAsia'),u'黑体')
            paragraph2_5.paragraph_format.space_after = Pt(0)
            paragraph2_5.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            #表2
            column_width = [1.5,11.0,2.22,1.81]
            device_num = len(err_result_great5)+1
            tablez2_5 = document.add_table(device_num,len(column_width),style='Table Grid')
            tablez2_5.alignment = WD_TAB_ALIGNMENT.CENTER
            for i in range(len(err_result_great5)+1):
                #print(err_temp.iloc[i-1]['turbine'])
                # if i == 0:
                #     print('biaoti')
                # else:
                #     print(err_temp.iloc[i-1]['turbine'])
                for j in range(len(column_width)):
                    if i == 0:
                        if j == 0:
                            fill_table(tablez2_5, x=i, y=j, content='机位号', font_name='黑体', font_size=10, bold=False, paragraph_alignment=WD_PARAGRAPH_ALIGNMENT.CENTER, startendmargin=0, width=column_width[j]) 
                        elif j == 1:
                            fill_table(tablez2_5, x=i, y=j, content='偏航误差角度-发电性能散点图', font_name='黑体', font_size=10, bold=False, paragraph_alignment=WD_PARAGRAPH_ALIGNMENT.CENTER, startendmargin=0, width=column_width[j]) 
                        elif j == 2:
                            fill_table(tablez2_5, x=i, y=j, content='偏航误差角(°)', font_name='黑体', font_size=10, bold=False, paragraph_alignment=WD_PARAGRAPH_ALIGNMENT.CENTER, startendmargin=0, width=column_width[j]) 
                        elif j == 3:
                            fill_table(tablez2_5, x=i, y=j, content='预估损失电量', font_name='黑体', font_size=10, bold=False, paragraph_alignment=WD_PARAGRAPH_ALIGNMENT.CENTER, startendmargin=0, width=column_width[j]) 
                    else:
                        # print(err_temp.iloc[i-1]['turbine'])
                        if j == 0:
                            fill_table(tablez2_5, x=i, y=j, content=conver_to_str(err_result_great5[i-1]['turbine']), font_name='Times New Roman', font_size=10, bold=False, paragraph_alignment=WD_PARAGRAPH_ALIGNMENT.CENTER, startendmargin=0, width=column_width[j]) 
                        elif j == 1:
                            figure_path = err_result_great5[i-1]['picture']#glob.glob(path+'/'+conver_to_str(err_result_great5[i-1]['turbine'])+'*yawerror.png')[0]
                            #print(figure_path)
                            fill_table(tablez2_5, x=i, y=j, picture=figure_path, paragraph_alignment=WD_PARAGRAPH_ALIGNMENT.CENTER, startendmargin=0, topbottommargin=0.01, width=column_width[j], picture_width=column_width[j]*0.8) 
                        elif j == 2:
                            fill_table(tablez2_5, x=i, y=j, content='{:.2f}'.format(err_result_great5[i-1]['yawerr']), font_name='Times New Roman', font_size=10, bold=False, paragraph_alignment=WD_PARAGRAPH_ALIGNMENT.CENTER, startendmargin=0, width=column_width[j]) 
                        elif j == 3:
                            fill_table(tablez2_5, x=i, y=j, content='{:.2%}'.format(err_result_great5[i-1]['loss']), font_name='Times New Roman', font_size=10, bold=False, paragraph_alignment=WD_PARAGRAPH_ALIGNMENT.CENTER, startendmargin=0, width=column_width[j]) 
        else:            
            # err_temp = err_result_all[(np.abs(err_result_all['yawerr'])<=5.0)]
            if len(err_result_less5) > 0:
                paragraph2_1_text = '分析时段内'+str(typeNameList[typei])+'机型各机组不存在较明显的偏航对风不正情况。如下图为某机组的偏航误差角度-发电性能散点图，误差角度小于5度。'      
                paragraph2_1 = document.add_paragraph()
                paragraph2_1_run = paragraph2_1.add_run(paragraph2_1_text)
                paragraph2_1_run.font.size = Pt(12)
                paragraph2_1_run.font.name = '仿宋'
                paragraph2_1_run._element.rPr.rFonts.set(qn('w:eastAsia'),u'仿宋')
                paragraph2_1.paragraph_format.first_line_indent = paragraph_run.font.size * 2
                paragraph2_1.paragraph_format.line_spacing = Pt(25)
                paragraph2_1.paragraph_format.space_after = Pt(6)
                paragraph2_1.paragraph_format.space_before = Pt(6)
                
                figure_path = err_result_less5[0]['picture']#glob.glob(path+'/'+str(err_temp.iloc[0]['turbine'])+'*yawerror.png')[0]
                paragraph_picture2_4 = document.add_paragraph()
                paragraph_picture2_4.alignment = WD_ALIGN_PARAGRAPH.CENTER
                paragraph_picture2_4_run = paragraph_picture2_4.add_run()
                paragraph_picture2_4_run.add_picture(figure_path,width=Cm(14))
                paragraph_picture2_4.paragraph_format.space_after = Pt(0)
                #图示
                picture_num += 1
                paragraph2_4_text = str('图'+str(picture_num)+' '+str(typeNameList[typei])+'机型'+str(err_temp.iloc[0]['turbine'])+'机组偏航误差角度-发电性能散点图')
                paragraph2_4 = document.add_paragraph()
                paragraph2_4_run = paragraph2_4.add_run(paragraph2_4_text)
                paragraph2_4_run.font.size = Pt(10)
                paragraph2_4_run.font.name = '黑体'
                paragraph2_4_run._element.rPr.rFonts.set(qn('w:eastAsia'),u'黑体')
                paragraph2_4.paragraph_format.space_after = Pt(16)
                paragraph2_4.alignment = WD_ALIGN_PARAGRAPH.CENTER
            else:
                paragraph2_1_text = '分析时段内'+str(typeNameList[typei])+'机型各机组不存在较明显的偏航对风不正情况。'      
                paragraph2_1 = document.add_paragraph()
                paragraph2_1_run = paragraph2_1.add_run(paragraph2_1_text)
                paragraph2_1_run.font.size = Pt(12)
                paragraph2_1_run.font.name = '仿宋'
                paragraph2_1_run._element.rPr.rFonts.set(qn('w:eastAsia'),u'仿宋')
                paragraph2_1.paragraph_format.first_line_indent = paragraph_run.font.size * 2
                paragraph2_1.paragraph_format.line_spacing = Pt(25)
                paragraph2_1.paragraph_format.space_after = Pt(6)
                paragraph2_1.paragraph_format.space_before = Pt(6)
#---------------------------------------------------------------------------------------

    for typei in range(len(metric_tongjiList)):#trubine_type
            
            # path = str(str(path_farm)+'/'+str(np.unique(Turbine_attr['turbineTypeID'])[typei]))
            turbine_err_all = []#pd.read_csv(str(path+'/turbine_err_all.csv'),header=[0],index_col=[0])
            for turbineName in farmInfo['wtid'][typeNameList[typei]]:
                resultBiasDirection = selectNavigationBiasControlPicture(farmInfo['farm_name'], typeNameList[typei], turbineName, startTime, endTime)
                if resultBiasDirection is not None:
                    dictBiasControl = {}
                    dictBiasControl['wtid'] = resultBiasDirection[0]
                    dictBiasControl['yaw_leiji_err'] = resultBiasDirection[2]
                    dictBiasControl['picture'] = resultBiasDirection[1]
                    if np.abs(dictBiasControl['yaw_leiji_err']) > 0: # and np.abs(dictBiasDirection['yawerr']) < 999
                        turbine_err_all.append(dictBiasControl)
            err_temp = turbine_err_all#[(np.abs(turbine_err_all['yaw_leiji_err'])>0)]
            if len(err_temp)>0:
                paragraph2_1_text = ('风力发电机组的偏航对风角度一般控制在-10°～+10°之间。分析时段内'+str(typeNameList[typei])+'机型部分机组偏航对风控制误差较大，详细见下表，偏航控制具有一定优化空间。')   
                paragraph2_1 = document.add_paragraph()
                paragraph2_1_run = paragraph2_1.add_run(paragraph2_1_text)
                paragraph2_1_run.font.size = Pt(12)
                paragraph2_1_run.font.name = '仿宋'
                paragraph2_1_run._element.rPr.rFonts.set(qn('w:eastAsia'),u'仿宋')
                paragraph2_1.paragraph_format.first_line_indent = paragraph_run.font.size * 2
                paragraph2_1.paragraph_format.line_spacing = Pt(25)
                paragraph2_1.paragraph_format.space_after = Pt(6)
                paragraph2_1.paragraph_format.space_before = Pt(6)
                
                #表例
                table_num = table_num + 1
                paragraph2_5_text = str('表'+str(table_num)+' '+str(typeNameList[typei])+'机型较大偏航控制误差统计表')
                paragraph2_5 = document.add_paragraph()
                paragraph2_5_run = paragraph2_5.add_run(paragraph2_5_text)
                paragraph2_5_run.font.size = Pt(10)
                paragraph2_5_run.font.name = '黑体'
                paragraph2_5_run._element.rPr.rFonts.set(qn('w:eastAsia'),u'黑体')
                paragraph2_5.paragraph_format.space_after = Pt(0)
                paragraph2_5.alignment = WD_ALIGN_PARAGRAPH.CENTER
                
                #表2
                column_width = [1.5,11.0]
                device_num = len(err_temp)+1
                tablez2_5 = document.add_table(device_num,len(column_width),style='Table Grid')
                tablez2_5.alignment = WD_TAB_ALIGNMENT.CENTER
                for i in range(device_num):
                    for j in range(len(column_width)):
                        if i == 0:
                            if j == 0:
                                fill_table(tablez2_5, x=i, y=j, content='机位号', font_name='黑体', font_size=10, bold=False, paragraph_alignment=WD_PARAGRAPH_ALIGNMENT.CENTER, startendmargin=0, width=column_width[j]) 
                            elif j == 1:
                                fill_table(tablez2_5, x=i, y=j, content='机组偏航控制角度直方图', font_name='黑体', font_size=10, bold=False, paragraph_alignment=WD_PARAGRAPH_ALIGNMENT.CENTER, startendmargin=0, width=column_width[j]) 
                        else:
                            if j == 0:
                                fill_table(tablez2_5, x=i, y=j, content=conver_to_str(err_temp[i-1]['wtid']), font_name='Times New Roman', font_size=10, bold=False, paragraph_alignment=WD_PARAGRAPH_ALIGNMENT.CENTER, startendmargin=0, width=column_width[j]) 
                            elif j == 1:
                                figure_path = err_temp[i-1]['picture']#glob.glob(path+'/'+str(err_temp.iloc[i-1]['wtid'])+'_wdir0.png')[0]
                                fill_table(tablez2_5, x=i, y=j, picture=figure_path, paragraph_alignment=WD_PARAGRAPH_ALIGNMENT.CENTER, startendmargin=0, topbottommargin=0.01, width=column_width[j], picture_width=column_width[j]*0.8) 
            else:
                pass
    #---------------------------------------------------------------------------------------
    paragraph2_2_text = '5.2. 变桨控制分析 '
    paragraph2_2 = document.add_paragraph()
    paragraph2_2_run = paragraph2_2.add_run(paragraph2_2_text)
    paragraph2_2_run.font.size = Pt(12)
    paragraph2_2_run.font.name = '黑体'
    paragraph2_2_run._element.rPr.rFonts.set(qn('w:eastAsia'),u'黑体')
    # paragraph2_2.paragraph_format.first_line_indent = paragraph_run.font.size * 2
    paragraph2_2.paragraph_format.space_before = Pt(12) 
    paragraph2_2.paragraph_format.space_after = Pt(0) 
    
    paragraph2_1_text = ('常见的制影响机组发电性能和载荷的变桨控问题包括最小桨距角异常、桨距角不平衡、变桨提前动作等。')   
    paragraph2_1 = document.add_paragraph()
    paragraph2_1_run = paragraph2_1.add_run(paragraph2_1_text)
    paragraph2_1_run.font.size = Pt(12)
    paragraph2_1_run.font.name = '仿宋'
    paragraph2_1_run._element.rPr.rFonts.set(qn('w:eastAsia'),u'仿宋')
    paragraph2_1.paragraph_format.first_line_indent = paragraph_run.font.size * 2
    paragraph2_1.paragraph_format.line_spacing = Pt(25)
    paragraph2_1.paragraph_format.space_after = Pt(0)
    paragraph2_1.paragraph_format.space_before = Pt(6)
    
    
    ###########最小桨距角异常
    for typei in range(len(metric_tongjiList)):#trubine_type
        
        # path = str(str(path_farm)+'/'+str(np.unique(Turbine_attr['turbineTypeID'])[typei]))
        turbine_err_all = []#pd.read_csv(str(path+'/turbine_err_all.csv'),header=[0],index_col=[0])
        for turbineName in farmInfo['wtid'][typeNameList[typei]]:
            resultMinPitch = selectPitchAnglePicture(farmInfo['farm_name'], typeNameList[typei], turbineName, startTime, endTime)
            if resultMinPitch is not None:
                dictMinPitch = {}
                dictMinPitch['wtid'] = resultMinPitch[0]
                dictMinPitch['pitch_min_err'] = resultMinPitch[3]
                dictMinPitch['picture'] = resultMinPitch[1]
                dictMinPitch['picture_compare'] = resultMinPitch[2]
                turbine_err_all.append(dictMinPitch)
        err_temp = turbine_err_all#[(np.abs(turbine_err_all['pitch_min_err'])>0)]
        if len(err_temp)>0:
            paragraph2_1_text = ('最小桨距角一般在0°左右，不同机型虽然需要根据叶片翼型及控制要求调整最小桨距角，但不会偏离0°太多。分析时段内'+str(typeNameList[typei])+
                                 '机型部分机组最小桨距角有异常，详细见下表，表中同时给出了异常机组与其邻近正常机组的功率曲线对比。')   
            paragraph2_1 = document.add_paragraph()
            paragraph2_1_run = paragraph2_1.add_run(paragraph2_1_text)
            paragraph2_1_run.font.size = Pt(12)
            paragraph2_1_run.font.name = '仿宋'
            paragraph2_1_run._element.rPr.rFonts.set(qn('w:eastAsia'),u'仿宋')
            paragraph2_1.paragraph_format.first_line_indent = paragraph_run.font.size * 2
            paragraph2_1.paragraph_format.line_spacing = Pt(25)
            paragraph2_1.paragraph_format.space_after = Pt(6)
            paragraph2_1.paragraph_format.space_before = Pt(6)
            
            #表例
            table_num = table_num + 1
            paragraph2_5_text = str('表'+str(table_num)+' '+str(typeNameList[typei])+'机型最小桨距角异常分析统计表')
            paragraph2_5 = document.add_paragraph()
            paragraph2_5_run = paragraph2_5.add_run(paragraph2_5_text)
            paragraph2_5_run.font.size = Pt(10)
            paragraph2_5_run.font.name = '黑体'
            paragraph2_5_run._element.rPr.rFonts.set(qn('w:eastAsia'),u'黑体')
            paragraph2_5.paragraph_format.space_after = Pt(0)
            paragraph2_5.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            #表2
            column_width = [1.5,6.5,6.5,1.51]
            device_num = len(err_temp)+1
            tablez2_5 = document.add_table(device_num,len(column_width),style='Table Grid')
            tablez2_5.alignment = WD_TAB_ALIGNMENT.CENTER
            for i in range(device_num):
                for j in range(len(column_width)):
                    if i == 0:
                        if j == 0:
                            fill_table(tablez2_5, x=i, y=j, content='机位号', font_name='黑体', font_size=10, bold=False, paragraph_alignment=WD_PARAGRAPH_ALIGNMENT.CENTER, startendmargin=0, width=column_width[j]) 
                        elif j == 1:
                            fill_table(tablez2_5, x=i, y=j, content='风速-桨距角散点图', font_name='黑体', font_size=10, bold=False, paragraph_alignment=WD_PARAGRAPH_ALIGNMENT.CENTER, startendmargin=0, width=column_width[j]) 
                        elif j == 2:
                            fill_table(tablez2_5, x=i, y=j, content='与邻近机组功率曲线对比图', font_name='黑体', font_size=10, bold=False, paragraph_alignment=WD_PARAGRAPH_ALIGNMENT.CENTER, startendmargin=0, width=column_width[j]) 
                        elif j == 3:
                            fill_table(tablez2_5, x=i, y=j, content='预估损失电量', font_name='黑体', font_size=10, bold=False, paragraph_alignment=WD_PARAGRAPH_ALIGNMENT.CENTER, startendmargin=0, width=column_width[j]) 
                    else:
                        if j == 0:
                            fill_table(tablez2_5, x=i, y=j, content=conver_to_str(err_temp[i-1]['wtid']), font_name='Times New Roman', font_size=10, bold=False, paragraph_alignment=WD_PARAGRAPH_ALIGNMENT.CENTER, startendmargin=0, width=column_width[j]) 
                        elif j == 1:
                            figure_path = err_temp[i-1]['picture']#glob.glob(path+'/'+conver_to_str(err_temp.iloc[i-1]['wtid'])+'_最小桨距角异常.png')[0]
                            fill_table(tablez2_5, x=i, y=j, picture=figure_path, paragraph_alignment=WD_PARAGRAPH_ALIGNMENT.CENTER, startendmargin=0, topbottommargin=0.01, width=column_width[j], picture_width=column_width[j],picture_height=5.5) 
                        elif j == 2:
                            figure_path = err_temp[i-1]['picture_compare']#glob.glob(path+'/'+conver_to_str(err_temp.iloc[i-1]['wtid'])+'_pitch_min_err.png')[0]
                            fill_table(tablez2_5, x=i, y=j, picture=figure_path, paragraph_alignment=WD_PARAGRAPH_ALIGNMENT.CENTER, startendmargin=0, topbottommargin=0.01, width=column_width[j], picture_width=column_width[j],picture_height=5.5) 
                        elif j == 3:
                            fill_table(tablez2_5, x=i, y=j, content='{:.2%}'.format(eval(err_temp[i-1]['pitch_min_loss'])), font_name='Times New Roman', font_size=10, bold=False, paragraph_alignment=WD_PARAGRAPH_ALIGNMENT.CENTER, startendmargin=0, width=column_width[j]) 
        else:
            paragraph2_1_text = '分析时段内'+str(typeNameList[typei])+'机型各机组不存在最小桨距角异常情况。'      
            paragraph2_1 = document.add_paragraph()
            paragraph2_1_run = paragraph2_1.add_run(paragraph2_1_text)
            paragraph2_1_run.font.size = Pt(12)
            paragraph2_1_run.font.name = '仿宋'
            paragraph2_1_run._element.rPr.rFonts.set(qn('w:eastAsia'),u'仿宋')
            paragraph2_1.paragraph_format.first_line_indent = paragraph_run.font.size * 2
            paragraph2_1.paragraph_format.line_spacing = Pt(25)
            paragraph2_1.paragraph_format.space_after = Pt(6)
            paragraph2_1.paragraph_format.space_before = Pt(6)
    
            
    ########变桨控制异常
    for typei in range(len(metric_tongjiList)):#trubine_type
        
        # path = str(str(path_farm)+'/'+str(np.unique(Turbine_attr['turbineTypeID'])[typei]))
        turbine_err_all = []#pd.read_csv(str(path+'/turbine_err_all.csv'),header=[0],index_col=[0])
        for turbineName in farmInfo['wtid'][typeNameList[typei]]:
            resultActionPitch = selectPitchActionPicture(farmInfo['farm_name'], typeNameList[typei], turbineName, startTime, endTime)
            if resultActionPitch is not None:
                dictActionPitch = {}
                dictActionPitch['wtid'] = resultActionPitch[0]
                dictActionPitch['picture'] = resultActionPitch[1]
                turbine_err_all.append(dictActionPitch)
        err_temp = turbine_err_all#[(np.abs(turbine_err_all['pitch_control_err'])>0)]
        if len(err_temp)>0:
            paragraph2_1_text = ('变桨控制一般在风机功率将近额定时才会动作，控制上出于防止机组超速、降低机组载荷、保证净空要求的考虑会适当的让变桨提前动作，'+
                                 '但变桨动作时风机实发功率一般不低于额定功率太多且变桨角度相对较小。额定功率以下过早或过大的变桨会影响风能吸收。'+
                                 '分析时段内'+str(typeNameList[typei])+
                                 '在功率****kw附件就开始变桨，且变桨角度较大，额定功率时桨距角已达到***度左右，下表中为各机组的功率-桨距角散点图。')   
            paragraph2_1 = document.add_paragraph()
            paragraph2_1_run = paragraph2_1.add_run(paragraph2_1_text)
            paragraph2_1_run.font.size = Pt(12)
            paragraph2_1_run.font.name = '仿宋'
            paragraph2_1_run._element.rPr.rFonts.set(qn('w:eastAsia'),u'仿宋')
            paragraph2_1.paragraph_format.first_line_indent = paragraph_run.font.size * 2
            paragraph2_1.paragraph_format.line_spacing = Pt(25)
            paragraph2_1.paragraph_format.space_after = Pt(6)
            paragraph2_1.paragraph_format.space_before = Pt(6)
            
            #表例
            table_num = table_num + 1
            paragraph2_5_text = str('表'+str(table_num)+' '+str(typeNameList[typei])+'机型异常机组功率-桨距角散点图')
            paragraph2_5 = document.add_paragraph()
            paragraph2_5_run = paragraph2_5.add_run(paragraph2_5_text)
            paragraph2_5_run.font.size = Pt(10)
            paragraph2_5_run.font.name = '黑体'
            paragraph2_5_run._element.rPr.rFonts.set(qn('w:eastAsia'),u'黑体')
            paragraph2_5.paragraph_format.space_after = Pt(0)
            paragraph2_5.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            #表2
            column_width = [1.5,12]
            device_num = len(err_temp)+1
            tablez2_5 = document.add_table(device_num,len(column_width),style='Table Grid')
            tablez2_5.alignment = WD_TAB_ALIGNMENT.CENTER
            for i in range(device_num):
                for j in range(len(column_width)):
                    if i == 0:
                        if j == 0:
                            fill_table(tablez2_5, x=i, y=j, content='机位号', font_name='黑体', font_size=10, bold=False, paragraph_alignment=WD_PARAGRAPH_ALIGNMENT.CENTER, startendmargin=0, width=column_width[j]) 
                        elif j == 1:
                            fill_table(tablez2_5, x=i, y=j, content='功率-桨距角散点图', font_name='黑体', font_size=10, bold=False, paragraph_alignment=WD_PARAGRAPH_ALIGNMENT.CENTER, startendmargin=0, width=column_width[j]) 
                    else:
                        if j == 0:
                            fill_table(tablez2_5, x=i, y=j, content=conver_to_str(err_temp[i-1]['wtid']), font_name='Times New Roman', font_size=10, bold=False, paragraph_alignment=WD_PARAGRAPH_ALIGNMENT.CENTER, startendmargin=0, width=column_width[j]) 
                        elif j == 1:                        
                            figure_path = err_temp[i-1]['picture']#glob.glob(path+'/'+str(err_temp.iloc[i-1]['wtid'])+'_pitch_control_err.png')[0]
                            fill_table(tablez2_5, x=i, y=j, picture=figure_path, paragraph_alignment=WD_PARAGRAPH_ALIGNMENT.CENTER, startendmargin=0, topbottommargin=0.01, width=column_width[j], picture_width=column_width[j]*0.8) 

        else:
            paragraph2_1_text = '分析时段内'+str(typeNameList[typei])+'机型各机组不存在变桨控制异常情况。'      
            paragraph2_1 = document.add_paragraph()
            paragraph2_1_run = paragraph2_1.add_run(paragraph2_1_text)
            paragraph2_1_run.font.size = Pt(12)
            paragraph2_1_run.font.name = '仿宋'
            paragraph2_1_run._element.rPr.rFonts.set(qn('w:eastAsia'),u'仿宋')
            paragraph2_1.paragraph_format.first_line_indent = paragraph_run.font.size * 2
            paragraph2_1.paragraph_format.line_spacing = Pt(25)
            paragraph2_1.paragraph_format.space_after = Pt(6)
            paragraph2_1.paragraph_format.space_before = Pt(6)
            
    
    ########变桨不平衡
    for typei in range(len(metric_tongjiList)):#trubine_type
        
        # path = str(str(path_farm)+'/'+str(np.unique(Turbine_attr['turbineTypeID'])[typei]))
        turbine_err_all = []#pd.read_csv(str(path+'/turbine_err_all.csv'),header=[0],index_col=[0])
        for turbineName in farmInfo['wtid'][typeNameList[typei]]:
            resultUnbalancePitch = selectPitchUnbalancePicture(farmInfo['farm_name'], typeNameList[typei], turbineName, startTime, endTime)
            if resultUnbalancePitch is not None:
                dictUnbalancePitch = {}
                dictUnbalancePitch['wtid'] = resultUnbalancePitch[0]
                dictUnbalancePitch['picture'] = resultUnbalancePitch[1]
                turbine_err_all.append(dictUnbalancePitch)
        err_temp = turbine_err_all#[(np.abs(turbine_err_all['pitch_balance_err'])>0)]
        if len(err_temp)>0:
            paragraph2_1_text = ('通常情况下风力发电机组三叶片变桨角度基本一致，即使采用独立变桨控制三叶片变桨角度也不会相差太大，'+
                                 '如果风机存在明显的三叶片变桨角度不一致情况，将可能导致机组气动在这不平衡、风机振动加剧、载荷增加等。'+
                                 '分析时段内'+str(typeNameList[typei])+
                                 '部分机组存在较明显的变桨角度不一致情况，下表中机组风速-桨距角散点图（绿色、黄色、红色分别表示不同叶片的变桨角度）。其余机组未见异常。')   
            paragraph2_1 = document.add_paragraph()
            paragraph2_1_run = paragraph2_1.add_run(paragraph2_1_text)
            paragraph2_1_run.font.size = Pt(12)
            paragraph2_1_run.font.name = '仿宋'
            paragraph2_1_run._element.rPr.rFonts.set(qn('w:eastAsia'),u'仿宋')
            paragraph2_1.paragraph_format.first_line_indent = paragraph_run.font.size * 2
            paragraph2_1.paragraph_format.line_spacing = Pt(25)
            paragraph2_1.paragraph_format.space_after = Pt(6)
            paragraph2_1.paragraph_format.space_before = Pt(6)
            
            #表例
            table_num = table_num + 1
            paragraph2_5_text = str('表'+str(table_num)+' '+str(typeNameList[typei])+'机型异常机组风速-桨距角散点图')
            paragraph2_5 = document.add_paragraph()
            paragraph2_5_run = paragraph2_5.add_run(paragraph2_5_text)
            paragraph2_5_run.font.size = Pt(10)
            paragraph2_5_run.font.name = '黑体'
            paragraph2_5_run._element.rPr.rFonts.set(qn('w:eastAsia'),u'黑体')
            paragraph2_5.paragraph_format.space_after = Pt(0)
            paragraph2_5.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            #表2
            column_width = [1.5,12]
            device_num = len(err_temp)+1
            tablez2_5 = document.add_table(device_num,len(column_width),style='Table Grid')
            tablez2_5.alignment = WD_TAB_ALIGNMENT.CENTER
            for i in range(device_num):
                for j in range(len(column_width)):
                    if i == 0:
                        if j == 0:
                            fill_table(tablez2_5, x=i, y=j, content='机位号', font_name='黑体', font_size=10, bold=False, paragraph_alignment=WD_PARAGRAPH_ALIGNMENT.CENTER, startendmargin=0, width=column_width[j]) 
                        elif j == 1:
                            fill_table(tablez2_5, x=i, y=j, content='风速-桨距角散点图', font_name='黑体', font_size=10, bold=False, paragraph_alignment=WD_PARAGRAPH_ALIGNMENT.CENTER, startendmargin=0, width=column_width[j]) 
                    else:
                        if j == 0:
                            fill_table(tablez2_5, x=i, y=j, content=conver_to_str(err_temp[i-1]['wtid']), font_name='Times New Roman', font_size=10, bold=False, paragraph_alignment=WD_PARAGRAPH_ALIGNMENT.CENTER, startendmargin=0, width=column_width[j]) 
                        elif j == 1:
                            figure_path = err_temp[i-1]['picture']#glob.glob(path+'/'+conver_to_str(err_temp.iloc[i-1]['wtid'])+'_pitch_balance_err.png')[0]
                            fill_table(tablez2_5, x=i, y=j, picture=figure_path, paragraph_alignment=WD_PARAGRAPH_ALIGNMENT.CENTER, startendmargin=0, topbottommargin=0.01, width=column_width[j], picture_width=column_width[j]*0.8) 
        else:
            paragraph2_1_text = '分析时段内'+str(typeNameList[typei])+'机型各机组不存在变桨角度不平衡情况。'      
            paragraph2_1 = document.add_paragraph()
            paragraph2_1_run = paragraph2_1.add_run(paragraph2_1_text)
            paragraph2_1_run.font.size = Pt(12)
            paragraph2_1_run.font.name = '仿宋'
            paragraph2_1_run._element.rPr.rFonts.set(qn('w:eastAsia'),u'仿宋')
            paragraph2_1.paragraph_format.first_line_indent = paragraph_run.font.size * 2
            paragraph2_1.paragraph_format.line_spacing = Pt(25)
            paragraph2_1.paragraph_format.space_after = Pt(6)
            paragraph2_1.paragraph_format.space_before = Pt(6)
            
    
            
    paragraph2_2_text = '5.3. 转矩控制分析 '
    paragraph2_2 = document.add_paragraph()
    paragraph2_2_run = paragraph2_2.add_run(paragraph2_2_text)
    paragraph2_2_run.font.size = Pt(12)
    paragraph2_2_run.font.name = '黑体'
    paragraph2_2_run._element.rPr.rFonts.set(qn('w:eastAsia'),u'黑体')
    # paragraph2_2.paragraph_format.first_line_indent = paragraph_run.font.size * 2
    paragraph2_2.paragraph_format.space_before = Pt(12) 
    paragraph2_2.paragraph_format.space_after = Pt(0) 
    
    paragraph2_1_text = ('通过机组运行数据分析转矩控制状况，在风机额定转速以下，为实现最大风能捕获，通过调节转矩控制转速跟随风速变化，确保机组风能利用系数(Cp)最大，此时功率与转速的三次方成正比。'+
                         '转矩控制异常将导致机组无法实现最大风能捕获，影响机组发电量，同时有可能导致转速振荡产生较大的疲劳载荷。'+
                         '产生转矩控制异常的原因通常为：1)风机控制策略异常；2)转矩控制参数异常。')   
    paragraph2_1 = document.add_paragraph()
    paragraph2_1_run = paragraph2_1.add_run(paragraph2_1_text)
    paragraph2_1_run.font.size = Pt(12)
    paragraph2_1_run.font.name = '仿宋'
    paragraph2_1_run._element.rPr.rFonts.set(qn('w:eastAsia'),u'仿宋')
    paragraph2_1.paragraph_format.first_line_indent = paragraph_run.font.size * 2
    paragraph2_1.paragraph_format.line_spacing = Pt(25)
    paragraph2_1.paragraph_format.space_after = Pt(0)
    paragraph2_1.paragraph_format.space_before = Pt(6)
    
    
    for typei in range(len(metric_tongjiList)):#trubine_type
        
        # path = str(str(path_farm)+'/'+str(np.unique(Turbine_attr['turbineTypeID'])[typei]))
        turbine_err_all = []#pd.read_csv(str(path+'/turbine_err_all.csv'),header=[0],index_col=[0])
        for turbineName in farmInfo['wtid'][typeNameList[typei]]:
            resultControlAction = selectTorqueControlPicture(farmInfo['farm_name'], typeNameList[typei], turbineName, startTime, endTime)
            if resultControlAction is not None:
                dictControlAction = {}
                dictControlAction['wtid'] = resultControlAction[0]
                dictControlAction['picture'] = resultControlAction[1]
                turbine_err_all.append(dictControlAction)
        err_temp = turbine_err_all#turbine_err_all[(np.abs(turbine_err_all['torque_kopt_err'])>0)]
        if len(err_temp)>0:
            paragraph2_1_text = ('分析时段内'+str(typeNameList[typei])+
                                 '部分机组存转矩控制异常，下图表中异常机组的转速-功率散点图，部分散点（红色）明显偏离最佳Cp控制区域。')   
            paragraph2_1 = document.add_paragraph()
            paragraph2_1_run = paragraph2_1.add_run(paragraph2_1_text)
            paragraph2_1_run.font.size = Pt(12)
            paragraph2_1_run.font.name = '仿宋'
            paragraph2_1_run._element.rPr.rFonts.set(qn('w:eastAsia'),u'仿宋')
            paragraph2_1.paragraph_format.first_line_indent = paragraph_run.font.size * 2
            paragraph2_1.paragraph_format.line_spacing = Pt(25)
            paragraph2_1.paragraph_format.space_after = Pt(6)
            paragraph2_1.paragraph_format.space_before = Pt(6)
            
            #表例
            table_num = table_num + 1
            paragraph2_5_text = str('表'+str(table_num)+' '+str(typeNameList[typei])+'机型异常机组转速-功率散点图')
            paragraph2_5 = document.add_paragraph()
            paragraph2_5_run = paragraph2_5.add_run(paragraph2_5_text)
            paragraph2_5_run.font.size = Pt(10)
            paragraph2_5_run.font.name = '黑体'
            paragraph2_5_run._element.rPr.rFonts.set(qn('w:eastAsia'),u'黑体')
            paragraph2_5.paragraph_format.space_after = Pt(0)
            paragraph2_5.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            #表2
            column_width = [1.5,12]
            device_num = len(err_temp)+1
            tablez2_5 = document.add_table(device_num,len(column_width),style='Table Grid')
            tablez2_5.alignment = WD_TAB_ALIGNMENT.CENTER
            for i in range(device_num):
                for j in range(len(column_width)):
                    if i == 0:
                        if j == 0:
                            fill_table(tablez2_5, x=i, y=j, content='机位号', font_name='黑体', font_size=10, bold=False, paragraph_alignment=WD_PARAGRAPH_ALIGNMENT.CENTER, startendmargin=0, width=column_width[j]) 
                        elif j == 1:
                            fill_table(tablez2_5, x=i, y=j, content='转速-功率散点图', font_name='黑体', font_size=10, bold=False, paragraph_alignment=WD_PARAGRAPH_ALIGNMENT.CENTER, startendmargin=0, width=column_width[j]) 
                    else:
                        if j == 0:
                            fill_table(tablez2_5, x=i, y=j, content=conver_to_str(err_temp[i-1]['wtid']), font_name='Times New Roman', font_size=10, bold=False, paragraph_alignment=WD_PARAGRAPH_ALIGNMENT.CENTER, startendmargin=0, width=column_width[j]) 
                        elif j == 1:
                            figure_path = err_temp[i-1]['picture']#glob.glob(path+'/'+conver_to_str(err_temp.iloc[i-1]['wtid'])+'_转矩kopt控制异常.png')[0]
                            fill_table(tablez2_5, x=i, y=j, picture=figure_path, paragraph_alignment=WD_PARAGRAPH_ALIGNMENT.CENTER, startendmargin=0, topbottommargin=0.01, width=column_width[j], picture_width=column_width[j]*0.8) 
        else:
            paragraph2_1_text = '分析时段内'+str(typeNameList[typei])+'机型各机组不存在转矩控制异常情况。'      #如下图为某机组的转速-功率散点图。
            paragraph2_1 = document.add_paragraph()
            paragraph2_1_run = paragraph2_1.add_run(paragraph2_1_text)
            paragraph2_1_run.font.size = Pt(12)
            paragraph2_1_run.font.name = '仿宋'
            paragraph2_1_run._element.rPr.rFonts.set(qn('w:eastAsia'),u'仿宋')
            paragraph2_1.paragraph_format.first_line_indent = paragraph_run.font.size * 2
            paragraph2_1.paragraph_format.line_spacing = Pt(25)
            paragraph2_1.paragraph_format.space_after = Pt(6)
            paragraph2_1.paragraph_format.space_before = Pt(6)
            
            # err_temp = turbine_err_all[(np.abs(turbine_err_all['torque_kopt_err'])<=0)]
            # wtid_random = np.random.randint(0,len(err_temp))
            # figure_path = glob.glob(path+'/'+(conver_to_str(err_temp.iloc[wtid_random]['wtid']))+'_分段1.png')[0]
            # paragraph_picture2_4 = document.add_paragraph()
            # paragraph_picture2_4.alignment = WD_ALIGN_PARAGRAPH.CENTER
            # paragraph_picture2_4_run = paragraph_picture2_4.add_run()
            # paragraph_picture2_4_run.add_picture(figure_path,width=Cm(14))
            # paragraph_picture2_4.paragraph_format.space_after = Pt(0)

            # #图示
            # picture_num += 1
            # paragraph2_4_text = str('图'+str(picture_num)+' '+str(np.unique(Turbine_attr['turbineTypeID'])[typei])+'机型'+str(err_temp.iloc[wtid_random]['wtid'])+'机组转速-功率散点图')
            # paragraph2_4 = document.add_paragraph()
            # paragraph2_4_run = paragraph2_4.add_run(paragraph2_4_text)
            # paragraph2_4_run.font.size = Pt(10)
            # paragraph2_4_run.font.name = '黑体'
            # paragraph2_4_run._element.rPr.rFonts.set(qn('w:eastAsia'),u'黑体')
            # paragraph2_4.paragraph_format.space_after = Pt(16)
            # paragraph2_4.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
    #-------------------------------------------------------------

    paragraph2_2_text = '5.4. 大部件异常分析 '
    paragraph2_2 = document.add_paragraph()
    paragraph2_2_run = paragraph2_2.add_run(paragraph2_2_text)
    paragraph2_2_run.font.size = Pt(12)
    paragraph2_2_run.font.name = '黑体'
    paragraph2_2_run._element.rPr.rFonts.set(qn('w:eastAsia'),u'黑体')
    # paragraph2_2.paragraph_format.first_line_indent = paragraph_run.font.size * 2
    paragraph2_2.paragraph_format.space_before = Pt(12) 
    paragraph2_2.paragraph_format.space_after = Pt(0) 
    
    paragraph2_1_text = ('分析发电机、齿轮箱、主轴、变流器、变桨电机、偏航电机、控制柜等部件的运行温度数据, 对比相邻两机组的对应部件的功率-温度图或者转速-温度图是否发生明显运行温度的异常偏移。')   
    paragraph2_1 = document.add_paragraph()
    paragraph2_1_run = paragraph2_1.add_run(paragraph2_1_text)
    paragraph2_1_run.font.size = Pt(12)
    paragraph2_1_run.font.name = '仿宋'
    paragraph2_1_run._element.rPr.rFonts.set(qn('w:eastAsia'),u'仿宋')
    paragraph2_1.paragraph_format.first_line_indent = paragraph_run.font.size * 2
    paragraph2_1.paragraph_format.line_spacing = Pt(25)
    paragraph2_1.paragraph_format.space_after = Pt(0)
    paragraph2_1.paragraph_format.space_before = Pt(6)
    
    
    for typei in range(len(metric_tongjiList)):#trubine_type
        
        # path = str(str(path_farm)+'/'+str(np.unique(Turbine_attr['turbineTypeID'])[typei]))
        turbine_err_all = []#pd.read_csv(str(path+'/turbine_err_all.csv'),header=[0],index_col=[0])
        for turbineName in farmInfo['wtid'][typeNameList[typei]]:
            resultDeviceTempture = selectDevicePicture(farmInfo['farm_name'], typeNameList[typei], turbineName, startTime, endTime)
            if resultDeviceTempture is not None:
                dictDeviceTempture = {}
                dictDeviceTempture['wtid'] = resultDeviceTempture[0]
                dictDeviceTempture['device'] = resultDeviceTempture[1]
                dictDeviceTempture['picture'] = resultDeviceTempture[2]
                turbine_err_all.append(dictDeviceTempture)
                err_temp = turbine_err_all#[(np.abs(turbine_err_all['pitch_balance_err'])>0)]
        if len(err_temp)>0:
            paragraph2_1_text = ('通常情况下相邻机组的部件的运行温度工况基本一致，'+
                                 +
                                 '分析时段内'+str(typeNameList[typei])+
                                 '部分机组存在较明显的运行温度异常，下表中展示机组功率-温度散点图或者转速-温度散点图（绿色、黄色、红色分别表示不同机组）。其余机组未见异常。')   
            paragraph2_1 = document.add_paragraph()
            paragraph2_1_run = paragraph2_1.add_run(paragraph2_1_text)
            paragraph2_1_run.font.size = Pt(12)
            paragraph2_1_run.font.name = '仿宋'
            paragraph2_1_run._element.rPr.rFonts.set(qn('w:eastAsia'),u'仿宋')
            paragraph2_1.paragraph_format.first_line_indent = paragraph_run.font.size * 2
            paragraph2_1.paragraph_format.line_spacing = Pt(25)
            paragraph2_1.paragraph_format.space_after = Pt(6)
            paragraph2_1.paragraph_format.space_before = Pt(6)
            
            #表例
            table_num = table_num + 1
            paragraph2_5_text = str('表'+str(table_num)+' '+str(typeNameList[typei])+'机型异常机组功率-温度散点图或者转速-温度散点图')
            paragraph2_5 = document.add_paragraph()
            paragraph2_5_run = paragraph2_5.add_run(paragraph2_5_text)
            paragraph2_5_run.font.size = Pt(10)
            paragraph2_5_run.font.name = '黑体'
            paragraph2_5_run._element.rPr.rFonts.set(qn('w:eastAsia'),u'黑体')
            paragraph2_5.paragraph_format.space_after = Pt(0)
            paragraph2_5.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            #表2
            column_width = [1.5,12]
            device_num = len(err_temp)+1
            tablez2_5 = document.add_table(device_num,len(column_width),style='Table Grid')
            tablez2_5.alignment = WD_TAB_ALIGNMENT.CENTER
            for i in range(device_num):
                for j in range(len(column_width)):
                    if i == 0:
                        if j == 0:
                            fill_table(tablez2_5, x=i, y=j, content='机位号', font_name='黑体', font_size=10, bold=False, paragraph_alignment=WD_PARAGRAPH_ALIGNMENT.CENTER, startendmargin=0, width=column_width[j]) 
                        elif j == 1:
                            fill_table(tablez2_5, x=i, y=j, content='功率-温度散点图或者转速-温度散点图', font_name='黑体', font_size=10, bold=False, paragraph_alignment=WD_PARAGRAPH_ALIGNMENT.CENTER, startendmargin=0, width=column_width[j]) 
                    else:
                        if j == 0:
                            fill_table(tablez2_5, x=i, y=j, content=conver_to_str(err_temp[i-1]['wtid']), font_name='Times New Roman', font_size=10, bold=False, paragraph_alignment=WD_PARAGRAPH_ALIGNMENT.CENTER, startendmargin=0, width=column_width[j]) 
                        elif j == 1:
                            figure_path = err_temp[i-1]['picture']#glob.glob(path+'/'+conver_to_str(err_temp.iloc[i-1]['wtid'])+'_pitch_balance_err.png')[0]
                            fill_table(tablez2_5, x=i, y=j, picture=figure_path, paragraph_alignment=WD_PARAGRAPH_ALIGNMENT.CENTER, startendmargin=0, topbottommargin=0.01, width=column_width[j], picture_width=column_width[j]*0.8) 
        else:
            paragraph2_1_text = '分析时段内'+str(typeNameList[typei])+'机型各机组不存在大部件温度异常情况。'      
            paragraph2_1 = document.add_paragraph()
            paragraph2_1_run = paragraph2_1.add_run(paragraph2_1_text)
            paragraph2_1_run.font.size = Pt(12)
            paragraph2_1_run.font.name = '仿宋'
            paragraph2_1_run._element.rPr.rFonts.set(qn('w:eastAsia'),u'仿宋')
            paragraph2_1.paragraph_format.first_line_indent = paragraph_run.font.size * 2
            paragraph2_1.paragraph_format.line_spacing = Pt(25)
            paragraph2_1.paragraph_format.space_after = Pt(6)
            paragraph2_1.paragraph_format.space_before = Pt(6)

    #-------------------------------------------------------------        
    paragraph2_2_text = '5.5. 其它 '
    paragraph2_2 = document.add_paragraph()
    paragraph2_2_run = paragraph2_2.add_run(paragraph2_2_text)
    paragraph2_2_run.font.size = Pt(12)
    paragraph2_2_run.font.name = '黑体'
    paragraph2_2_run._element.rPr.rFonts.set(qn('w:eastAsia'),u'黑体')
    # paragraph2_2.paragraph_format.first_line_indent = paragraph_run.font.size * 2
    paragraph2_2.paragraph_format.space_before = Pt(12) 
    paragraph2_2.paragraph_format.space_after = Pt(0) 
    
    #小节5
    paragraph2 = document.add_paragraph(style='ListNumber')
    paragraph2_run = paragraph2.add_run('总结与建议')
    paragraph2_run.font.size = Pt(14)
    paragraph2_run.font.name = '黑体'
    paragraph2_run._element.rPr.rFonts.set(qn('w:eastAsia'),u'黑体')
    paragraph2.paragraph_format.space_before = Pt(12)
    paragraph2.paragraph_format.space_after = Pt(6)

    file_word = farmInfo['path_farm'] + '/'+ datetime.strftime(execute_time, "%Y-%m-%d_%H-%M-%S")+"_" + windfarm_name + '能效评估报告.docx'
    file_word = os.path.abspath(file_word)
    document.save(file_word)
    removeFile(farmInfo['path_farm'])

    return file_word #os.path.abspath(farmInfo['path_farm'] + '/' + datetime.strftime(execute_time, "%Y-%m-%d_%H-%M-%S")+"_"+ windfarm_name + '能效评估报告.docx')





def removeFile(path):
    # 设置截止日期，例如，删除10天之前的文件
    cutoff_time = datetime.now() - timedelta(days=10)
    
    # 指定要检查的目录
    directory = path
    #扩展名
    ext = ["docx", "doc"]
    
    # 遍历目录中的所有文件
    for filename in os.listdir(directory):
        file_path = os.path.join(directory, filename)
        file_extension = file_path.split(".")[-1]
        if file_extension in ext:
            try:
                # 获取文件的修改时间
                file_mtime = os.path.getmtime(file_path)
                # 将时间转换为datetime对象进行比较
                file_time = datetime.fromtimestamp(file_mtime)
                # 如果文件修改时间早于截止时间，则删除文件
                if file_time < cutoff_time:
                    os.remove(file_path)
                    print(f"Deleted: {filename}")
            except Exception as e:
                errorInfomation = traceback.format_exc()
                print(f"#########################Error processing {filename}#######################################3:\n {errorInfomation}")

    '''
    


##########################################################################################

    #小节4
    paragraph3 = document.add_paragraph(style='ListNumber')
    paragraph3_run = paragraph3.add_run('控制性能分析')
    paragraph3_run.font.size = Pt(12)
    paragraph3_run.font.name = '黑体'
    paragraph3_run._element.rPr.rFonts.set(qn('w:eastAsia'),u'黑体')
    paragraph3.paragraph_format.space_before = Pt(12)

    #文字
    paragraph3_1_text = '从产业数据中台提取了各机组的转速、桨距角、功率、风速、风向等测点数据，从偏航控制、变桨控制、转矩控制等方面分析控制策略对机组发电性能的影响。'
    paragraph3_1 = document.add_paragraph()
    paragraph3_1_run = paragraph3_1.add_run(paragraph3_1_text)
    paragraph3_1_run.font.size = Pt(12)
    paragraph3_1_run.font.name = '仿宋'
    paragraph3_1_run._element.rPr.rFonts.set(qn('w:eastAsia'),u'仿宋')
    paragraph3_1.paragraph_format.first_line_indent = paragraph_run.font.size * 2
    paragraph3_1.paragraph_format.space_after = Pt(6)

    #子小节
    paragraph3_2_text = '4.1.偏航控制分析'
    paragraph3_2 = document.add_paragraph()
    paragraph3_2_run = paragraph3_2.add_run(paragraph3_2_text)
    paragraph3_2_run.font.size = Pt(11)
    paragraph3_2_run.font.name = '黑体'
    paragraph3_2_run._element.rPr.rFonts.set(qn('w:eastAsia'),u'黑体')
    # paragra3h15_2.paragraph_format.first_line_indent = paragraph_run.font.size * 2
    paragraph3_2.paragraph_format.space_after = Pt(6) 

    #文字
    paragraph3_3_text = '偏航控制存在两方面的问题，分别是偏航控制误差较大和偏航对风不正。'
    paragraph3_3 = document.add_paragraph()
    paragraph3_3_run = paragraph3_3.add_run(paragraph3_3_text)
    paragraph3_3_run.font.size = Pt(12)
    paragraph3_3_run.font.name = '仿宋'
    paragraph3_3_run._element.rPr.rFonts.set(qn('w:eastAsia'),u'仿宋')
    paragraph3_3.paragraph_format.first_line_indent = paragraph_run.font.size * 2
    paragraph3_3.paragraph_format.space_after = Pt(6)
    #文字
    paragraph3_4_text = '风场大部分机组的误差范围在-20°～+20°之间，相比其它风场控制误差较大，见下图，偏航控制具有一定优化空间。'
    paragraph3_4 = document.add_paragraph()
    paragraph3_4_run = paragraph3_4.add_run(paragraph3_4_text)
    paragraph3_4_run.font.size = Pt(12)
    paragraph3_4_run.font.name = '仿宋'
    paragraph3_4_run._element.rPr.rFonts.set(qn('w:eastAsia'),u'仿宋')
    paragraph3_4.paragraph_format.first_line_indent = paragraph_run.font.size * 2
    paragraph3_4.paragraph_format.space_after = Pt(6)

    #表例
    table_num = table_num + 1
    paragraph3_5_text = str('表'+str(table_num)+ '偏航控制误差分析')
    paragraph3_5 = document.add_paragraph()
    paragraph3_5_run = paragraph3_5.add_run(paragraph3_5_text)
    paragraph3_5_run.font.size = Pt(10)
    paragraph3_5_run.font.name = '黑体'
    paragraph3_5_run._element.rPr.rFonts.set(qn('w:eastAsia'),u'黑体')
    paragraph3_5.paragraph_format.first_line_indent = paragraph_run.font.size * 2
    paragraph3_5.paragraph_format.space_after = Pt(0)
    paragraph3_5.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    #表6偏航
    device_num = table_wdir.shape[0] + 1
    column_width = [2.16,11.4]
    tablez3_5 = document.add_table(device_num,len(column_width),style='Table Grid')
    #tablez.17utofit = True
    tablez3_5.alignment = WD_TAB_ALIGNMENT.CENTER
    for i in range(device_num):
        for j in range(len(column_width)):
            if i == 0:
                if j == 0:
                   fill_table(tablez3_5, x=i, y=j, content='机位号', font_name='黑体', font_size=10, bold=True, paragraph_alignment=WD_PARAGRAPH_ALIGNMENT.CENTER, startendmargin=0, width=column_width[j]) 
                elif j == 1:
                   fill_table(tablez3_5, x=i, y=j, content='风速-偏航控制角度直方图', font_name='黑体', font_size=10, bold=True, paragraph_alignment=WD_PARAGRAPH_ALIGNMENT.CENTER, startendmargin=0, width=column_width[j]) 
                
                
            else:
                if j == 0:
                   fill_table(tablez3_5, x=i, y=j, content=str(table_wdir.iloc[i-1]['turbine_name']), font_name='仿宋', font_size=10, bold=False, paragraph_alignment=WD_PARAGRAPH_ALIGNMENT.CENTER, startendmargin=0, width=column_width[j]) 
                elif j == 1:
                   fill_table(tablez3_5, x=i, y=j, picture=table_wdir.iloc[i-1]['figure_name'], paragraph_alignment=WD_PARAGRAPH_ALIGNMENT.CENTER, startendmargin=0, topbottommargin=0.01, width=column_width[j], picture_width=column_width[j]*0.8) 

    if table_err_result_yawerr_greater_5.empty != True:
        #文字
        paragraph3_6_text = '风场'+str(table_err_result_yawerr_greater_5.shape[0])+'台机组存在5度以上的对风偏航误差，产生偏航误差的原因通常为：1)风向标标定不准；2)风向标安装松动；3)叶轮尾流影响；4)偏航控制策略。当偏航误差角度超过5度，将对机组发电量造成明显影响。各机组偏航误差统计详见下表，表中散点图为对风角度与发电性能散点图，机组的发电性能最高点明显偏离了偏航对风的0度位置，校正后预计可提升发电量'+'{:.2f}'.format(table_err_result_yawerr_greater_5['loss'].mean()*100)+'%。'
        paragraph3_6 = document.add_paragraph()
        paragraph3_6_run = paragraph3_6.add_run(paragraph3_6_text)
        paragraph3_6_run.font.size = Pt(12)
        paragraph3_6_run.font.name = '仿宋'
        paragraph3_6_run._element.rPr.rFonts.set(qn('w:eastAsia'),u'仿宋')
        paragraph3_6.paragraph_format.first_line_indent = paragraph_run.font.size * 2
        paragraph3_6.paragraph_format.space_after = Pt(6)
        paragraph3_6.paragraph_format.space_before = Pt(6)

        #表例
        table_num = table_num + 1
        paragraph3_7_text = str('表'+str(table_num)+ '偏航对风偏差统计表')
        paragraph3_7 = document.add_paragraph()
        paragraph3_7_run = paragraph3_7.add_run(paragraph3_7_text)
        paragraph3_7_run.font.size = Pt(10)
        paragraph3_7_run.font.name = '黑体'
        paragraph3_7_run._element.rPr.rFonts.set(qn('w:eastAsia'),u'黑体')
        paragraph3_7.paragraph_format.first_line_indent = paragraph_run.font.size * 2
        paragraph3_7.paragraph_format.space_after = Pt(0)
        paragraph3_7.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        #表7
        device_num = table_err_result_yawerr_greater_5.shape[0] + 1
        column_width = [2.15,10.25,2,1.5]
        tablez3_7 = document.add_table(device_num,len(column_width),style='Table Grid')
        #tablez.77utofit = True
        tablez3_7.alignment = WD_TAB_ALIGNMENT.CENTER
        for i in range(device_num):
            for j in range(len(column_width)):
                if i == 0:
                    if j == 0:
                        fill_table(tablez3_7, x=i, y=j, content='机位号', font_name='黑体', font_size=10, bold=True, paragraph_alignment=WD_PARAGRAPH_ALIGNMENT.CENTER, startendmargin=0, width=column_width[j]) 
                    elif j == 1:
                        fill_table(tablez3_7, x=i, y=j, content='图形分析', font_name='黑体', font_size=10, bold=True, paragraph_alignment=WD_PARAGRAPH_ALIGNMENT.CENTER, startendmargin=0, width=column_width[j]) 
                    elif j == 2:
                        fill_table(tablez3_7, x=i, y=j, content='偏航误差角(°)', font_name='黑体', font_size=10, bold=True, paragraph_alignment=WD_PARAGRAPH_ALIGNMENT.CENTER, startendmargin=0, width=column_width[j]) 
                    elif j == 3:
                        fill_table(tablez3_7, x=i, y=j, content='损失电量', font_name='黑体', font_size=10, bold=True, paragraph_alignment=WD_PARAGRAPH_ALIGNMENT.CENTER, startendmargin=0, width=column_width[j]) 
                    
                else:
                    if j == 0:
                        fill_table(tablez3_7, x=i, y=j, content=str(table_err_result_yawerr_greater_5.iloc[i-1]['turbine_name']), font_name='仿宋', font_size=10, bold=False, paragraph_alignment=WD_PARAGRAPH_ALIGNMENT.CENTER, startendmargin=0, width=column_width[j]) 
                    elif j == 1:
                        fill_table(tablez3_7, x=i, y=j, picture=table_err_result_yawerr_greater_5.iloc[i-1]['figure_name'], paragraph_alignment=WD_PARAGRAPH_ALIGNMENT.CENTER, startendmargin=0, topbottommargin=0.01, width=column_width[j], picture_width=column_width[j]*0.8) 
                    elif j == 2:
                        fill_table(tablez3_7, x=i, y=j, content='{:.2f}'.format(table_err_result_yawerr_greater_5.iloc[i-1]['yawerr']), font_name='仿宋', font_size=10, bold=False, paragraph_alignment=WD_PARAGRAPH_ALIGNMENT.CENTER, startendmargin=0, width=column_width[j]) 
                    elif j == 3:
                        fill_table(tablez3_7, x=i, y=j, content='{:.2f}'.format(table_err_result_yawerr_greater_5.iloc[i-1]['loss']*100)+'%', font_name='仿宋', font_size=10, bold=False, paragraph_alignment=WD_PARAGRAPH_ALIGNMENT.CENTER, startendmargin=0, width=column_width[j]) 

#-----------------------------------------------------------------------------------------
    #子小节
    paragraph3_8_text = '4.2.变桨控制分析'
    paragraph3_8 = document.add_paragraph()
    paragraph3_8_run = paragraph3_8.add_run(paragraph3_8_text)
    paragraph3_8_run.font.size = Pt(11)
    paragraph3_8_run.font.name = '黑体'
    paragraph3_8_run._element.rPr.rFonts.set(qn('w:eastAsia'),u'黑体')
    # paragra3h85_2.paragraph_format.first_line_indent = paragraph_run.font.size * 2
    paragraph3_8.paragraph_format.space_after = Pt(6) 
    paragraph3_8.paragraph_format.space_before = Pt(12)

    if table_wspd_yaw.shape[0] > 0:
        #文字
        yaw_turbines = ''
        for i in range(table_wspd_yaw.shape[0]):
            if yaw_turbines == '':
                yaw_turbines = yaw_turbines + table_wspd_yaw.iloc[i]['turbine_name']
            else:
                yaw_turbines = yaw_turbines + '，' + table_wspd_yaw.iloc[i]['turbine_name'] 

        paragraph3_9_text = '风机在风速较小时一般保持最佳桨距角运行，最佳桨距角由叶片翼型及控制要求而定，一般在0度附近。风场'+yaw_turbines+'机组最小桨距角设置有误，最小桨距角控制明显超出合理范围，导致机组实际功率曲线远低于其余机组的功率曲线，如下表所示。'
        paragraph3_9 = document.add_paragraph()
        paragraph3_9_run = paragraph3_9.add_run(paragraph3_9_text)
        paragraph3_9_run.font.size = Pt(12)
        paragraph3_9_run.font.name = '仿宋'
        paragraph3_9_run._element.rPr.rFonts.set(qn('w:eastAsia'),u'仿宋')
        paragraph3_9.paragraph_format.first_line_indent = paragraph_run.font.size * 2
        paragraph3_9.paragraph_format.space_after = Pt(6)

        #表例
        table_num = table_num + 1
        paragraph3_10_text = str('表'+str(table_num)+ '最佳桨距角异常分析')
        paragraph3_10 = document.add_paragraph()
        paragraph3_10_run = paragraph3_10.add_run(paragraph3_10_text)
        paragraph3_10_run.font.size = Pt(10)
        paragraph3_10_run.font.name = '黑体'
        paragraph3_10_run._element.rPr.rFonts.set(qn('w:eastAsia'),u'黑体')
        paragraph3_10.paragraph_format.first_line_indent = paragraph_run.font.size * 2
        paragraph3_10.paragraph_format.space_after = Pt(0)
        paragraph3_10.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        #表8
        device_num = table_wspd_yaw.shape[0] + 1
        column_width = [1.11,7.37,7.22,1.11]
        tablez3_10 = document.add_table(device_num,len(column_width),style='Table Grid')
        #tablez.107utofit = True
        tablez3_10.alignment = WD_TAB_ALIGNMENT.CENTER
        for i in range(device_num):
            for j in range(len(column_width)):
                if i == 0:
                    if j == 0:
                        fill_table(tablez3_10, x=i, y=j, content='机位号', font_name='黑体', font_size=10, bold=True, paragraph_alignment=WD_PARAGRAPH_ALIGNMENT.CENTER, startendmargin=0, width=column_width[j]) 
                    elif j == 1:
                        fill_table(tablez3_10, x=i, y=j, content='风速-桨距角散点图', font_name='黑体', font_size=10, bold=True, paragraph_alignment=WD_PARAGRAPH_ALIGNMENT.CENTER, startendmargin=0, width=column_width[j]) 
                    elif j == 2:
                        fill_table(tablez3_10, x=i, y=j, content='与邻近机组功率曲线对比', font_name='黑体', font_size=10, bold=True, paragraph_alignment=WD_PARAGRAPH_ALIGNMENT.CENTER, startendmargin=0, width=column_width[j]) 
                    elif j == 3:
                        fill_table(tablez3_10, x=i, y=j, content='损失电量', font_name='黑体', font_size=10, bold=True, paragraph_alignment=WD_PARAGRAPH_ALIGNMENT.CENTER, startendmargin=0, width=column_width[j]) 
                    
                else:
                    if j == 0:
                        fill_table(tablez3_10, x=i, y=j, content=str(table_wspd_yaw.iloc[i-1]['turbine_name']), font_name='仿宋', font_size=10, bold=False, paragraph_alignment=WD_PARAGRAPH_ALIGNMENT.CENTER, startendmargin=0, width=column_width[j]) 
                    elif j == 1:
                        fill_table(tablez3_10, x=i, y=j, picture=table_wspd_yaw.iloc[i-1]['figure_name_1'], paragraph_alignment=WD_PARAGRAPH_ALIGNMENT.CENTER, startendmargin=0, topbottommargin=0.01, width=column_width[j], picture_width=column_width[j]*0.8) 
                    elif j == 2:
                        fill_table(tablez3_10, x=i, y=j, picture=table_wspd_yaw.iloc[i-1]['figure_name_2'], paragraph_alignment=WD_PARAGRAPH_ALIGNMENT.CENTER, startendmargin=0, topbottommargin=0.01, width=column_width[j], picture_width=column_width[j]*0.8) 
                    elif j == 3:
                        fill_table(tablez3_10, x=i, y=j, content='{:.2f}'.format(table_wspd_yaw.iloc[i-1]['loss']*100)+'%', font_name='仿宋', font_size=10, bold=False, paragraph_alignment=WD_PARAGRAPH_ALIGNMENT.CENTER, startendmargin=0, width=column_width[j]) 
                    
    paragraph3_11_text = '变桨控制一般在风机功率将近额定时才会动作，控制上出于防止机组超速、降低机组载荷、保证净空要求的考虑会适当的让变桨提前动作，但变桨动作时功率一般不低于额定功率的80%且变桨角度相对较小。'+str(table_pwrat_pitch.iloc[0]['turbine_name'])+'等机组在远低于额定功率时就开始变桨，且变桨角度较大，影响额定转速段的风能吸收。'
    paragraph3_11 = document.add_paragraph()
    paragraph3_11_run = paragraph3_11.add_run(paragraph3_11_text)
    paragraph3_11_run.font.size = Pt(12)
    paragraph3_11_run.font.name = '仿宋'
    paragraph3_11_run._element.rPr.rFonts.set(qn('w:eastAsia'),u'仿宋')
    paragraph3_11.paragraph_format.first_line_indent = paragraph_run.font.size * 2
    paragraph3_11.paragraph_format.space_after = Pt(6)
    paragraph3_11.paragraph_format.space_before = Pt(6)

    #表例
    table_num = table_num + 1
    paragraph3_12_text = str('表'+str(table_num)+ '变桨控制分析')
    paragraph3_12 = document.add_paragraph()
    paragraph3_12_run = paragraph3_12.add_run(paragraph3_12_text)
    paragraph3_12_run.font.size = Pt(10)
    paragraph3_12_run.font.name = '黑体'
    paragraph3_12_run._element.rPr.rFonts.set(qn('w:eastAsia'),u'黑体')
    paragraph3_12.paragraph_format.first_line_indent = paragraph_run.font.size * 2
    paragraph3_12.paragraph_format.space_after = Pt(0)
    paragraph3_12.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    #表9
    device_num = table_pwrat_pitch.shape[0] + 1
    column_width = [2.2,12.2]
    tablez3_12 = document.add_table(device_num,len(column_width),style='Table Grid')
    #tablez.127utofit = True
    tablez3_12.alignment = WD_TAB_ALIGNMENT.CENTER
    for i in range(device_num):
        for j in range(len(column_width)):
            if i == 0:
                if j == 0:
                    fill_table(tablez3_12, x=i, y=j, content='机位号', font_name='黑体', font_size=10, bold=True, paragraph_alignment=WD_PARAGRAPH_ALIGNMENT.CENTER, startendmargin=0, width=column_width[j]) 
                elif j == 1:
                    fill_table(tablez3_12, x=i, y=j, content='散点图', font_name='黑体', font_size=10, bold=True, paragraph_alignment=WD_PARAGRAPH_ALIGNMENT.CENTER, startendmargin=0, width=column_width[j]) 
                
            else:
                if j == 0:
                    fill_table(tablez3_12, x=i, y=j, content=str(table_pwrat_pitch.iloc[i-1]['turbine_name']), font_name='仿宋', font_size=10, bold=False, paragraph_alignment=WD_PARAGRAPH_ALIGNMENT.CENTER, startendmargin=0, width=column_width[j]) 
                elif j == 1:
                    fill_table(tablez3_12, x=i, y=j, picture=table_pwrat_pitch.iloc[i-1]['figure_name'], paragraph_alignment=WD_PARAGRAPH_ALIGNMENT.CENTER, startendmargin=0, topbottommargin=0.01, width=column_width[j], picture_width=column_width[j]*0.8) 
 #-----------------------------------------------------------------------------------------                   

    #子小节
    paragraph3_13_text = '4.3.转矩控制分析'
    paragraph3_13 = document.add_paragraph()
    paragraph3_13_run = paragraph3_13.add_run(paragraph3_13_text)
    paragraph3_13_run.font.size = Pt(11)
    paragraph3_13_run.font.name = '黑体'
    paragraph3_13_run._element.rPr.rFonts.set(qn('w:eastAsia'),u'黑体')
    # paragra3h135_2.paragraph_format.first_line_indent = paragraph_run.font.size * 2
    paragraph3_13.paragraph_format.space_after = Pt(6)                
    paragraph3_13.paragraph_format.space_before = Pt(12)

    #文字
    if table_rotspd_pwrat.empty == False:
        paragraph3_13_text = '通过机组运行数据分析转矩控制状况，在风机额定转速以下，为实现最大风能捕获，通过调节转矩控制转速跟随风速变化，此时功率与转速的三次方成正比。风场'+str(table_rotspd_pwrat.iloc[0]['turbine_name'])+'等机组转速-功率散点明显发散，存在异常，如下表所示。'
    else:
        paragraph3_13_text = '通过机组运行数据分析转矩控制状况，在风机额定转速以下，为实现最大风能捕获，通过调节转矩控制转速跟随风速变化，此时功率与转速的三次方成正比。风场'+'没有机组转速-功率散点明显发散，不存在异常。'
    paragraph3_13 = document.add_paragraph()
    paragraph3_13_run = paragraph3_13.add_run(paragraph3_13_text)
    paragraph3_13_run.font.size = Pt(12)
    paragraph3_13_run.font.name = '仿宋'
    paragraph3_13_run._element.rPr.rFonts.set(qn('w:eastAsia'),u'仿宋')
    paragraph3_13.paragraph_format.first_line_indent = paragraph_run.font.size * 2
    paragraph3_13.paragraph_format.space_after = Pt(6)
    #文字
    paragraph3_14_text = '产生转矩控制异常的原因通常为：1)风机控制策略异常；2)转矩控制参数异常。'
    paragraph3_14 = document.add_paragraph()
    paragraph3_14_run = paragraph3_14.add_run(paragraph3_14_text)
    paragraph3_14_run.font.size = Pt(12)
    paragraph3_14_run.font.name = '仿宋'
    paragraph3_14_run._element.rPr.rFonts.set(qn('w:eastAsia'),u'仿宋')
    paragraph3_14.paragraph_format.first_line_indent = paragraph_run.font.size * 2
    paragraph3_14.paragraph_format.space_after = Pt(6)
    #文字
    paragraph3_15_text = '转矩控制异常将导致机组无法实现最大风能捕获，影响机组发电量，同时有可能导致转速振荡产生较大的疲劳载荷。建议联系整机厂家检查机组控制策略。'
    paragraph3_15 = document.add_paragraph()
    paragraph3_15_run = paragraph3_15.add_run(paragraph3_15_text)
    paragraph3_15_run.font.size = Pt(12)
    paragraph3_15_run.font.name = '仿宋'
    paragraph3_15_run._element.rPr.rFonts.set(qn('w:eastAsia'),u'仿宋')
    paragraph3_15.paragraph_format.first_line_indent = paragraph_run.font.size * 2
    paragraph3_15.paragraph_format.space_after = Pt(6)

    if table_rotspd_pwrat.empty == False:
        #表例
        table_num = table_num + 1
        paragraph3_16_text = str('表'+str(table_num)+ '转矩控制异常分析')
        paragraph3_16 = document.add_paragraph()
        paragraph3_16_run = paragraph3_16.add_run(paragraph3_16_text)
        paragraph3_16_run.font.size = Pt(10)
        paragraph3_16_run.font.name = '黑体'
        paragraph3_16_run._element.rPr.rFonts.set(qn('w:eastAsia'),u'黑体')
        paragraph3_16.paragraph_format.first_line_indent = paragraph_run.font.size * 2
        paragraph3_16.paragraph_format.space_after = Pt(0)
        paragraph3_16.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        #表10
        device_num = table_rotspd_pwrat.shape[0] + 1
        column_width = [2.2,12.2]
        tablez3_16 = document.add_table(device_num,len(column_width),style='Table Grid')
        #tablez.167utofit = True
        tablez3_16.alignment = WD_TAB_ALIGNMENT.CENTER
        for i in range(device_num):
            for j in range(len(column_width)):
                if i == 0:
                    if j == 0:
                        fill_table(tablez3_16, x=i, y=j, content='机位号', font_name='黑体', font_size=10, bold=True, paragraph_alignment=WD_PARAGRAPH_ALIGNMENT.CENTER, startendmargin=0, width=column_width[j]) 
                    elif j == 1:
                        fill_table(tablez3_16, x=i, y=j, content='散点图', font_name='黑体', font_size=10, bold=True, paragraph_alignment=WD_PARAGRAPH_ALIGNMENT.CENTER, startendmargin=0, width=column_width[j]) 
                    
                else:
                    if j == 0:
                        fill_table(tablez3_16, x=i, y=j, content=str(table_rotspd_pwrat.iloc[i-1]['turbine_name']), font_name='仿宋', font_size=10, bold=False, paragraph_alignment=WD_PARAGRAPH_ALIGNMENT.CENTER, startendmargin=0, width=column_width[j]) 
                    elif j == 1:
                        fill_table(tablez3_16, x=i, y=j, picture=table_rotspd_pwrat.iloc[i-1]['figure_name'], paragraph_alignment=WD_PARAGRAPH_ALIGNMENT.CENTER, startendmargin=0, topbottommargin=0.01, width=column_width[j], picture_width=column_width[j]*0.8) 

##########################################################################################

    #小节5
    paragraph4 = document.add_paragraph(style='ListNumber')
    paragraph4_run = paragraph4.add_run('总结与建议')
    paragraph4_run.font.size = Pt(12)
    paragraph4_run.font.name = '黑体'
    paragraph4_run._element.rPr.rFonts.set(qn('w:eastAsia'),u'黑体')
    paragraph4.paragraph_format.space_before = Pt(12)


   
##########################################################################################

    InsertPageNumber(document)
    
    paragraph = document.add_paragraph()
    run = paragraph.add_run()
    fldChar = OxmlElement('w:fldChar') # creates a new element
    fldChar.set(qn('w:fldCharType'), 'begin') # sets attribute on element
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve') # sets attribute on element
    instrText.text = 'TOC \\o "1-3" \\h \\z \\u' # change 1-3 depending on heading levels you need
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'separate')
    fldChar3 = OxmlElement('w:t')
    # fldChar3.text = "Right-click to update field."
    fldChar2.append(fldChar3)
    fldChar4 = OxmlElement('w:fldChar')
    fldChar4.set(qn('w:fldCharType'), 'end')
    r_element = run._r
    r_element.append(fldChar)
    r_element.append(instrText)
    r_element.append(fldChar2)
    r_element.append(fldChar4)
    p_element = paragraph._p
        
    document.save(path_save + '/' + windfarm_name + '检测报告UAT.docx')

    return path_save + '/' + windfarm_name + '检测报告UAT.docx'

'''



    